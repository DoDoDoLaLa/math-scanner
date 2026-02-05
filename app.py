
# -*- coding: utf-8 -*-
from __future__ import annotations

import io
import os
import re
import json
import time
import base64
import tempfile
import concurrent.futures
from functools import lru_cache
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any, Callable, Iterable

import streamlit as st
from streamlit.errors import StreamlitDuplicateElementKey


# --- tiny helpers ---
def _has_cjk(s: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", s or ""))

# --- Academic features: glossary + post-translation QA (auto-retry) ---
GLOSS_PLACEHOLDER_RE = re.compile(r"__GLOSS\d+__")
def _is_ascii_alnum_term(s: str) -> bool:
    t = (s or "").strip()
    return bool(t) and bool(re.fullmatch(r"[A-Za-z0-9]+", t))


def _is_cjk_term(s: str) -> bool:
    return bool(re.search(r"[一-鿿]", s or ""))


def _compile_glossary_pattern(src: str, *, allow_substring_match: bool) -> re.Pattern:
    esc = re.escape(src)
    if allow_substring_match:
        return re.compile(esc)
    if _is_ascii_alnum_term(src):
        return re.compile(rf"\b{esc}\b")
    if _is_cjk_term(src):
        return re.compile(esc)
    # Mixed/symbol term (e.g., machine learning, p-value): non-word boundaries via lookaround.
    return re.compile(rf"(?<![A-Za-z0-9_]){esc}(?![A-Za-z0-9_])")

def parse_glossary_mapping(text: str) -> Dict[str, str]:
    """Parse glossary lines into dict.
    Supported formats per line:
    - src => dst
    - src -> dst
    - src\t dst
    - src,dst   (only if exactly 2 columns)
    Blank lines / comment lines (# ...) are ignored.
    """
    mapping: Dict[str, str] = {}
    if not text:
        return mapping
    for raw in str(text).splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        if "\t" in line:
            parts = [p.strip() for p in line.split("\t") if p.strip()]
        elif "=>" in line:
            parts = [p.strip() for p in line.split("=>", 1)]
        elif "->" in line:
            parts = [p.strip() for p in line.split("->", 1)]
        else:
            parts = [p.strip() for p in line.split(",")]
        if len(parts) != 2 or not parts[0] or not parts[1]:
            continue
        mapping[parts[0]] = parts[1]
    return mapping

def apply_glossary_placeholders(
    text: str,
    glossary: Dict[str, str],
    *,
    allow_substring_match: bool = False,
) -> Tuple[str, Dict[str, str]]:
    """Replace glossary source terms with stable placeholders in one regex pass.

    Safety goals:
    - no destructive substring pollution (cat -> catastrophe, AI -> PAIN)
    - longest-first matching on overlaps
    - no nested / repeated placeholder replacement
    """
    if not text or not glossary:
        return text, {}

    entries = [(src, tgt) for src, tgt in glossary.items() if (src or "").strip()]
    if not entries:
        return text, {}

    # Longest-first for overlap stability. Keep insertion order as secondary key.
    entries = sorted(entries, key=lambda kv: len(kv[0]), reverse=True)

    patterns: List[Tuple[re.Pattern, str, str, str]] = []
    placeholder_to_target: Dict[str, str] = {}
    for i, (src, tgt) in enumerate(entries):
        ph = f"__GLOSS{i}__"
        patterns.append((_compile_glossary_pattern(src, allow_substring_match=allow_substring_match), src, ph, tgt))

    occupied = [False] * len(text)
    selected: List[Tuple[int, int, str, str]] = []  # (start, end, placeholder, target)
    for patt, _src, ph, tgt in patterns:
        for m in patt.finditer(text):
            a, b = m.span()
            if a >= b:
                continue
            if any(occupied[a:b]):
                continue
            selected.append((a, b, ph, tgt))
            for k in range(a, b):
                occupied[k] = True

    if not selected:
        return text, {}

    selected.sort(key=lambda x: x[0])
    out_parts: List[str] = []
    cur = 0
    for a, b, ph, tgt in selected:
        if a < cur:
            continue
        out_parts.append(text[cur:a])
        out_parts.append(ph)
        placeholder_to_target[ph] = tgt
        cur = b
    out_parts.append(text[cur:])
    return "".join(out_parts), placeholder_to_target

def restore_glossary_placeholders(text: str, placeholder_to_target: Dict[str, str]) -> str:
    if not text or not placeholder_to_target:
        return text
    out = text
    # Replace longer placeholders first (though they are uniform)
    for ph, tgt in sorted(placeholder_to_target.items(), key=lambda kv: len(kv[0]), reverse=True):
        out = out.replace(ph, tgt)
    return out

def qa_detect_issues(text: str, *, dst_lang: str, forbid_cjk_when_en: bool = True) -> List[str]:
    issues: List[str] = []
    t = text or ""
    # 1) Leftover CJK when target is English
    if forbid_cjk_when_en and dst_lang in ("English", "en"):
        if _has_cjk(t):
            issues.append("CJK_LEFT")
    # 2) Enforce $/$$ delimiters only (no \( \) \[ \] equation env) for this app's Pandoc stability
    if re.search(r"\\\(|\\\)|\\\[|\\\]", t):
        issues.append("MATH_DELIMS_NOT_DOLLARS")
    if re.search(r"\\begin\{equation\*?\}", t):
        issues.append("EQUATION_ENV_FOUND")
    # 3) Placeholder leaks (glossary)
    if GLOSS_PLACEHOLDER_RE.search(t):
        issues.append("GLOSSARY_PLACEHOLDER_LEAK")
    return issues

def llm_translate_with_qa(
    *,
    client: OpenAI,
    model: str,
    prompt: str,
    timeout_s: int,
    base_retries: int = 6,
    qa_retries: int = 2,
    dst_lang: str,
) -> str:
    """Call LLM, then run simple QA; if issues exist, auto-retry with a repair prompt."""
    last_text = ""
    for attempt in range(qa_retries + 1):
        res = safe_chat_completions(
            client=client,
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=8192,
            timeout_s=int(timeout_s),
            retries=int(base_retries),
        )
        if res.error_message:
            raise TranslateCallError(res.error_message, raw=None)
        text = (res.text or "").strip()
        last_text = text
        issues = qa_detect_issues(text, dst_lang=dst_lang)
        if not issues:
            return text
        # Build a minimal repair prompt: keep structure, keep formulas, remove forbidden patterns.
        prompt = (
            "你刚才的翻译输出存在质量问题，需要你在不改变含义/结构的前提下修复并重新输出。\n"
            "严格要求：\n"
            "- 只输出修复后的正文（不要解释，不要额外字段）。\n"
            "- 数学公式一律使用 $...$（行内）与 $$...$$（行间），不要使用 \\(..\\)、\\[..\\] 或 equation 环境。\n"
            "- 如果目标语言是 English/en，则输出中不得出现中文/CJK 字符。\n\n"
            f"问题：{', '.join(issues)}\n\n"
            "【原始提示】\n" + prompt + "\n\n"
            "【你上一次的输出】\n" + text + "\n\n"
            "请直接给出修复后的最终输出："
        )
    return last_text

from PIL import Image, ImageFilter, ImageOps

from openai import OpenAI


import httpx
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import pypandoc

# LaTeX-in-text -> Word equation (OMML) support (optional)
try:
    from lxml import etree  # type: ignore
    import latex2mathml.converter  # type: ignore
    HAVE_LATEX_OMML = True
except Exception:
    HAVE_LATEX_OMML = False

# PDF rendering (optional but recommended)
try:
    import fitz  # PyMuPDF
    HAVE_PYMUPDF = True
except Exception:
    HAVE_PYMUPDF = False


# ============================================================
# 0) App UI
# ============================================================
st.set_page_config(page_title="Ark OCR / DOCX Translate / LaTeX Export", layout="wide")

st.markdown(
    """<style>
    .stTabs [data-baseweb="tab"] {font-size: 15px; padding: 10px 14px;}
    .stTabs [aria-selected="true"] {font-weight: 700;}
    .hint-card {background: #f6f8fa; border: 1px solid #e5e7eb; padding: 12px 14px; border-radius: 10px;}
    </style>""",
    unsafe_allow_html=True,
)

st.title("学术 OCR & Word→LaTeX 工具（Ark EP 已接入）")

st.caption(
    "面向论文/讲义/教材：PDF/图片OCR→Markdown/Word；Word（含可编辑公式）→LaTeX/Markdown；Word 原排版内翻译（best-effort）。"
)

def ensure_pandoc():
    try:
        _ = pypandoc.get_pandoc_path()
    except OSError:
        try:
            pypandoc.download_pandoc()
        except Exception:
            pass

# ============================================================
# 0.5) Pipeline module (math conversions) + caching
# ============================================================
# 将 DOCX<->Pandoc<->DOCX(OMML) 的重活拆到 math_pipeline.py，app.py 主要负责 UI。
import math_pipeline as mp

# 统一读取上传文件字节：优先 getvalue()，避免 UploadedFile.read() 导致“二次读取为空”
def read_uploaded_bytes(uploaded_file) -> bytes:
    return mp.read_uploaded_bytes(uploaded_file)

# --- 缓存：避免同一文件反复跑 Pandoc / AI（省时间、省费用） ---
@st.cache_data(show_spinner=False, max_entries=32)
def _cached_docx_to_md_for_math(docx_sha: str, docx_bytes: bytes) -> str:
    return mp.docx_to_pandoc_markdown_for_math(docx_bytes, wrap_none=True)

@st.cache_data(show_spinner=False, max_entries=32)
def _cached_roundtrip_omml(docx_sha: str, docx_bytes: bytes) -> bytes:
    out, _md = mp.docx_roundtrip_make_equations_editable(docx_bytes)
    return out

@st.cache_data(show_spinner=False, max_entries=32)
def _cached_roundtrip_omml_with_md(docx_sha: str, docx_bytes: bytes):
    return mp.docx_roundtrip_make_equations_editable(docx_bytes)

# AI 结果也缓存（同文件 + 同参数 + 同模型）——避免重复付费
@st.cache_data(show_spinner=False, max_entries=16)
def _cached_ai_roundtrip(
    docx_sha: str,
    docx_bytes: bytes,
    model: str,
    max_batch_chars: int,
    timeout_s: int,
    out_tokens: int,
    base_url: str,
):
    client = get_ark_client(int(timeout_s))

    def call_llm(prompt: str) -> str:
        res = safe_chat_completions(
            client=client,
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=int(out_tokens),
            timeout_s=int(timeout_s),
            retries=6,
        )
        if res.error_message:
            raise RuntimeError(res.error_message)
        return res.text or ""

    # 缓存函数里不要用 st.progress 之类的 UI 组件（会破坏缓存）
    return mp.docx_ai_roundtrip_make_equations_editable(
        docx_bytes,
        call_llm=call_llm,
        max_batch_chars=int(max_batch_chars),
        glossary_map=parse_glossary_mapping(st.session_state.get("glossary_text", "")),
        enable_qa_retry=bool(st.session_state.get("qa_enable", True)),
        qa_retries=int(st.session_state.get("qa_retries", 2)),
        progress_cb=None,
    )






# --- 翻译增强：直接 DOCX -> LLM 输出 Markdown（OCR $$ 风格）---
# 目标：省去“先用 Pandoc 抽 Markdown 再翻译”的步骤。我们把 docx 正文按段落抽取成纯文本，
# 并把 Word 原生公式（OMML）以内嵌占位形式交给 LLM 转成 LaTeX（$...$ / $$...$$），然后再用 Pandoc 写回 Word(OMML)。
def _iter_paragraph_children_xml(p: Paragraph) -> List[str]:
    """Return paragraph content as a linear list of chunks, preserving the order of text runs and OMML nodes."""
    chunks: List[str] = []
    for child in list(p._p):
        # Word text run
        if child.tag.endswith("}r"):
            # collect all <w:t> under this run
            ts = child.findall(".//" + qn("w:t"))
            if ts:
                chunks.append("".join([(t.text or "") for t in ts]))
        # OMML math nodes (m:oMath / m:oMathPara)
        elif child.tag.endswith("}oMath") or child.tag.endswith("}oMathPara"):
            chunks.append({"__OMML__": child.xml, "__TAG__": child.tag}.get("__OMML__", child.xml))
        else:
            # ignore other nodes (bookmarks, etc.)
            pass
    return chunks

def _linearize_docx_with_omml_markers(docx_bytes: bytes) -> Tuple[str, int]:
    """DOCX bytes -> linearized text with OMML markers embedded. Returns (text, eq_count)."""
    doc = Document(io.BytesIO(docx_bytes))
    paras = iter_all_paragraphs_extended(doc)

    eq_idx = 0
    out_lines: List[str] = []

    for p in paras:
        parts = _iter_paragraph_children_xml(p)
        if not parts:
            continue
        buf = []
        for part in parts:
            if not part:
                continue
            # detect if this chunk is OMML xml (starts with <m:oMath or <m:oMathPara)
            s = part.strip()
            if s.startswith("<m:oMathPara") or s.startswith("<m:oMath"):
                display = 1 if s.startswith("<m:oMathPara") else 0
                xml_b64 = base64.b64encode(s.encode("utf-8")).decode("utf-8")
                marker = f"__OMML_EQ_{eq_idx}__{{display={display},xml_b64={xml_b64}}}__OMML_EQ_{eq_idx}__"
                buf.append(marker)
                eq_idx += 1
            else:
                buf.append(part)
        line = "".join(buf)
        if line.strip():
            out_lines.append(line)

    # keep paragraph boundaries as blank lines (markdown-ish)
    return "\n\n".join(out_lines).replace("\r\n", "\n").replace("\r", "\n"), eq_idx

def _chunk_text_for_llm(text: str, max_chars: int) -> List[str]:
    """Chunk long text by paragraph breaks to keep structure."""
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text:
        return []
    paras = text.split("\n\n")
    chunks: List[str] = []
    cur: List[str] = []
    cur_len = 0
    for p in paras:
        p = p.strip("\n")
        if not p:
            continue
        add = p + "\n\n"
        if cur and cur_len + len(add) > max_chars:
            chunks.append("".join(cur).strip())
            cur, cur_len = [], 0
        cur.append(add)
        cur_len += len(add)
    if cur:
        chunks.append("".join(cur).strip())
    return chunks

def _translate_docx_direct_to_markdown_ocr_style(
    docx_bytes: bytes,
    *,
    client: OpenAI,
    model: str,
    src_lang: str,
    dst_lang: str,
    timeout_s: int,
    max_batch_chars: int,
    glossary_map: Optional[Dict[str, str]] = None,
    glossary_allow_substring: bool = False,
    enable_qa_retry: bool = True,
    qa_retries: int = 2,
    progress_cb: Optional[Callable[[int, int], None]] = None,
    attempt_msg_cb: Optional[Callable[[str], None]] = None,
) -> Tuple[str, Dict[str, Any]]:
    """Directly translate DOCX content to OCR-style Markdown via LLM (no Pandoc pre-conversion)."""
    linear, eq_cnt = _linearize_docx_with_omml_markers(docx_bytes)
    chunks = _chunk_text_for_llm(linear, max_chars=int(max_batch_chars))
    stats: Dict[str, Any] = {"chunks": len(chunks), "eq_count": int(eq_cnt), "cjk_left": 0}

    if not chunks:
        return "", stats

    out_parts: List[str] = []
    total = len(chunks)

    base_prompt = DOCX_TRANSLATE_OCR_STYLE_PROMPT.replace("__SRC_LANG__", src_lang).replace("__DST_LANG__", dst_lang)

    for i, ch in enumerate(chunks, start=1):
        if progress_cb:
            progress_cb(i, total)
        if attempt_msg_cb:
            attempt_msg_cb(f"直接翻译批次 {i}/{total}（包含 OMML 公式标记，LLM 将转为 $/$$）")

        ch2, ph_map = apply_glossary_placeholders(
            ch,
            glossary_map or {},
            allow_substring_match=bool(glossary_allow_substring),
        )
        prompt = base_prompt + "\n\n" + ch2

        if enable_qa_retry:
            text = llm_translate_with_qa(
                client=client,
                model=model,
                prompt=prompt,
                timeout_s=int(timeout_s),
                base_retries=6,
                qa_retries=int(qa_retries),
                dst_lang=dst_lang,
            )
        else:
            res = safe_chat_completions(
                client=client,
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0,
                max_tokens=8192,
                timeout_s=int(timeout_s),
                retries=6,
            )
            if res.error_message:
                raise TranslateCallError(res.error_message, raw=None)
            text = (res.text or "").strip()

        text = restore_glossary_placeholders(text, ph_map)
        out_parts.append(text.strip())

    md = mp.normalize_md("\n\n".join([p for p in out_parts if p.strip()]))

    # Post-fix: ensure math delimiters are dollars and stable for Pandoc
    md = mp.sanitize_tex_math_for_pandoc(md)
    md = mp.normalize_pandoc_math(md, display_style="dollars", inline_style="dollars")
    md = mp.normalize_md(md)

    if dst_lang in ("English", "en"):
        stats["cjk_left"] = int(sum(1 for line in md.splitlines() if _has_cjk(line)))
    return md, stats
# --- 翻译增强：Pandoc 中转（OCR $$ 风格） ---
def _split_md_blocks_preserve_spacing(md: str):
    """
    Split markdown into blocks separated by blank lines, but keep the separators so we can re-join
    without changing layout too much.
    Returns list of (is_sep, text).
    """
    parts = re.split(r"(\n\s*\n+)", md.replace("\r\n", "\n").replace("\r", "\n"))
    out = []
    for p in parts:
        if p is None or p == "":
            continue
        if re.fullmatch(r"\n\s*\n+", p):
            out.append((True, p))
        else:
            out.append((False, p))
    return out

def _translate_markdown_ocr_style(
    md: str,
    *,
    client: OpenAI,
    model: str,
    src_lang: str,
    dst_lang: str,
    timeout_s: int,
    max_batch_chars: int,
    progress_cb: Optional[Callable[[int, int], None]] = None,
    attempt_msg_cb: Optional[Callable[[str], None]] = None,
) -> Tuple[str, Dict[str, Any]]:
    """
    Translate markdown via Ark translate model, while preserving math in $...$/$$...$$.
    This route is closer to Tab① OCR output: formulas are kept in dollars delimiters,
    and text is translated block-by-block.

    Returns: (translated_md, stats)
    """
    md = mp.normalize_md(md)
    blocks = _split_md_blocks_preserve_spacing(md)

    items = []
    refs = []  # (block_idx, mapping)
    bid = 0
    for i, (is_sep, text) in enumerate(blocks):
        if is_sep:
            continue
        t = text.strip("\n")
        if not t.strip():
            continue
        protected, mapping = protect_math(t)
        bid += 1
        item_id = f"b{bid}"
        items.append({"id": item_id, "segments": [protected]})
        refs.append((item_id, i, mapping))

    stats = {"blocks_total": len([1 for is_sep,_ in blocks if not is_sep]), "blocks_translated": len(items), "blocks_untranslated_after": 0}

    if not items:
        return md, stats

    # batch translate with adaptive split (robust for long docs)
    batches = chunk_items_for_api(items, max_chars=max_batch_chars)
    total = len(batches)
    translated_map_all: Dict[str, List[str]] = {}

    for bi, batch in enumerate(batches, start=1):
        if progress_cb:
            progress_cb(bi, total)
        if attempt_msg_cb:
            attempt_msg_cb(f"Pandoc 翻译批次 {bi}/{total}：items={len(batch)}")

        translated_map = llm_translate_items_json(
            client=client,
            model=model,
            items=batch,
            src_lang=src_lang,
            dst_lang=dst_lang,
            timeout_s=timeout_s,
            retries=6,
            on_attempt=(lambda a, n, ph: attempt_msg_cb(f"Pandoc 批次 {bi}/{total}：尝试 {a}/{n} · {ph}") if attempt_msg_cb else None),
        )
        translated_map_all.update(translated_map)

    # write back blocks
    for item_id, block_idx, mapping in refs:
        seg = translated_map_all.get(item_id, [""])[0]
        seg = restore_tokens(seg, mapping)
        blocks[block_idx] = (False, seg)

    translated_md = mp.normalize_md("".join([t for _, t in blocks]))

    # Stats: count residue CJK when target is English
    if dst_lang in ("English", "en"):
        left = sum(1 for _, t in blocks if not _has_cjk(t) is False and _has_cjk(t))
        stats["blocks_untranslated_after"] = int(left)

    # Always re-sanitize for Pandoc math stability
    translated_md = mp.sanitize_tex_math_for_pandoc(translated_md)
    translated_md = mp.normalize_pandoc_math(translated_md, display_style="dollars", inline_style="dollars")
    translated_md = mp.normalize_md(translated_md)
    return translated_md, stats


@st.cache_data(show_spinner=False, max_entries=16)
def _cached_pandoc_translate_ocr_route(
    docx_sha: str,
    docx_bytes: bytes,
    src_lang: str,
    dst_lang: str,
    timeout_s: int,
    max_batch_chars: int,
    glossary_text: str = "",
    glossary_allow_substring: bool = False,
) -> Tuple[str, bytes, str, Dict[str, Any]]:
    """
    Direct route (no Pandoc pre-conversion):
      - DOCX -> LLM 直接输出 Markdown（OCR $$ 风格，数学用 $...$ / $$...$$）
      - Pandoc: Markdown -> DOCX（写回为可编辑 OMML）
      - Pandoc: Markdown -> LaTeX
    """
    client = get_ark_client(int(timeout_s))
    md_tr, stats = _translate_docx_direct_to_markdown_ocr_style(
        docx_bytes,
        client=client,
        model=get_default_model(),
        src_lang=src_lang,
        dst_lang=dst_lang,
        timeout_s=int(timeout_s),
        max_batch_chars=int(max_batch_chars),
        glossary_map=parse_glossary_mapping(glossary_text),
        progress_cb=None,
        attempt_msg_cb=None,
        glossary_allow_substring=bool(glossary_allow_substring),
    )
    out_docx = mp.pandoc_markdown_to_docx(md_tr, reference_docx_bytes=docx_bytes)
    latex = mp.pandoc_markdown_to_latex(md_tr)
    return md_tr, out_docx, latex, stats


ensure_pandoc()


# ============================================================
# 1) Prompts
# ============================================================
OCR_PROMPT_ZH = r"""
你是一个严谨的中文学术 OCR 转写器。请从图片中提取内容并输出 Markdown，要求：
1) 把图片中所有可见文字逐字转写（保持中文，不要翻译，不要改写）。
2) 数学公式必须转为 LaTeX，并尽量保持原位：
   - 行内用 $...$
   - 行间用 $$...$$
   - 如果有公式编号（如 (6)(7) 或 \tag{6}），保留编号信息。
3) 保持原本顺序，表格用 Markdown 表格输出。
4) 只输出 Markdown 正文，不要解释。
5) 保持原有阅读顺序、换行、项目符号、标题层级。
""".strip()

TRANSLATE_PROMPT_TEMPLATE = r"""
你是一个专业学术翻译引擎。请把以下内容从 __SRC_LANG__ 翻译到 __DST_LANG__。

严格要求：
- 输入是 JSON，包含 items 列表，每个 item 有 id 和 segments（字符串列表）。
- 输出必须是 JSON，结构必须为：{"items":[{"id":"...","segments":[...]}, ...]}
- segments 的数量必须与输入完全一致；每个 segments[i] 对应翻译输入 segments[i]。
- 保留所有占位符不变：例如 __MATH_0__、__KEEP_12__、{{ }} 这种标记必须原样输出，不能翻译、不能改大小写、不能删。
- LaTeX 代码、公式环境（如 \begin{equation}...\end{equation} 或 $...$）不得改动。
- 不要添加多余字段、不要输出解释、不要 Markdown。
""".strip()


# --- Direct DOCX -> Markdown (OCR $$ style) translation prompt ---
DOCX_TRANSLATE_OCR_STYLE_PROMPT = r"""
你是一个严谨的学术翻译与排版引擎。请把我给你的“Word 文档内容抽取文本”从 __SRC_LANG__ 翻译到 __DST_LANG__，
并直接输出 Markdown，要求尽量模仿本应用 Tab① 的 OCR 输出风格。

输入说明：
- 我会提供一段“抽取后的正文”，其中夹杂一些公式标记：__OMML_EQ_0__、__OMML_EQ_1__...（对应 Word 原生公式的 OMML XML，已用 base64 编码）。
- 公式标记会以如下形式出现：
  __OMML_EQ_k__{display=0|1,xml_b64=...}__OMML_EQ_k__
  你需要把它替换成 LaTeX 数学公式，并放回原位（行内/行间由 display 决定）。

严格要求：
1) 只输出 Markdown 正文，不要解释，不要标题，不要 JSON，不要代码块围栏。
2) 普通文本做学术翻译：语义准确、尽量保留原有结构（换行、列表、标题层级）。
3) 数学公式必须转换为 LaTeX，并统一分隔符：
   - 行内公式用 $...$
   - 行间公式用 $$...$$
   - 不要输出 \( ... \)、\[ ... \)、\begin{equation}...\end{equation}
4) 公式内容尽量忠实于 OMML：不要改写变量含义；不要引入新宏包命令；保留可能的编号/标签信息（若能识别）。
5) 保留占位符/保留字不变：例如 __MATH_0__、__KEEP_12__、{{ }} 之类若出现在正文中，必须原样保留。
6) 如果目标语言是 English/en：输出中不得出现中文（CJK 字符）。

下面是输入正文（含公式标记）：
""".strip()


FORMULA_OCR_TO_LATEX_PROMPT_ZH = r"""
你是数学公式 OCR 转写器。请只输出该图片中最核心公式对应的 LaTeX 代码，要求：
1) 只输出 LaTeX 公式文本，不要解释，不要代码块围栏。
2) 不要添加额外自然语言。
3) 默认输出可用于行间公式的主体（不强制包含 $$）。
4) 若有多个公式，按从上到下拼接，使用 \\ 分行。
""".strip()

SMART_GLOSSARY_EXTRACT_PROMPT_ZH = r"""
你是学术翻译术语抽取助手。请从给定文本中提取翻译时应保持一致的术语候选。

输出必须是 JSON 对象，格式：
{
  "items": [
    {"src":"", "suggested":"", "type":"term|person|place", "freq_estimate":1, "note":""}
  ]
}

规则：
1) 仅输出最重要 top-K 候选，避免常见普通词。
2) src 为原文术语，suggested 为建议译法（可与 src 相同），type 只能是 term/person/place。
3) freq_estimate 给出 1-100 的粗略频次估计。
4) note 可用于提示大小写、复数、词形变体（例如 AI/AIs, model/models）。
5) 只输出 JSON，不要解释。
""".strip()


def smart_extract_glossary_candidates(
    *,
    client: OpenAI,
    model: str,
    preview_text: str,
    src_lang: str,
    dst_lang: str,
    top_k: int,
    timeout_s: int,
) -> List[Dict[str, Any]]:
    if not preview_text.strip():
        return []
    prompt = (
        SMART_GLOSSARY_EXTRACT_PROMPT_ZH
        + f"\n\n源语言: {src_lang}\n目标语言: {dst_lang}\nTop-K: {int(top_k)}\n"
        + "\n文本如下：\n"
        + preview_text
    )
    res = safe_chat_completions(
        client=client,
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=4096,
        timeout_s=int(timeout_s),
        retries=6,
        response_format={"type": "json_object"},
    )
    if res.error_message:
        raise TranslateCallError(res.error_message, raw=None)
    obj = extract_json_object(res.text)
    items = obj.get("items") if isinstance(obj, dict) else None
    if not isinstance(items, list):
        return []
    out: List[Dict[str, Any]] = []
    for it in items:
        if not isinstance(it, dict):
            continue
        src = str(it.get("src", "")).strip()
        if not src:
            continue
        out.append({
            "src": src,
            "suggested": str(it.get("suggested", "")).strip(),
            "type": str(it.get("type", "term")).strip() or "term",
            "freq_estimate": int(it.get("freq_estimate", 1) or 1),
            "note": str(it.get("note", "")).strip(),
        })
    return out[: int(top_k)]


def llm_translate_items_json(
    *,
    client: OpenAI,
    model: str,
    items: List[dict],
    src_lang: str,
    dst_lang: str,
    timeout_s: int,
    retries: int = 6,
    on_attempt: Optional[Callable[[int, int, str], None]] = None,
) -> Dict[str, List[str]]:
    """Translate a batch of items via chat model with strict JSON IO.

    This is used ONLY by the 'Pandoc 中转（OCR $$ 风格）' route so we can:
    - force output to be Markdown (not extra prose)
    - keep math placeholders untouched
    - strongly discourage target-language drift (e.g., English output containing Chinese)
    """
    if not items:
        return {}

    prompt = TRANSLATE_PROMPT_TEMPLATE.replace("__SRC_LANG__", src_lang).replace("__DST_LANG__", dst_lang)

    # Extra constraints for this route:
    # - Output segments are plain Markdown text
    # - Keep all math as-is and keep using $...$ / $$...$$ (no \(\), \[\], equation env)
    prompt += (
        "\n\n额外约束（必须遵守）：\n"
        "- 你的每个 segments 输出都必须是 Markdown 正文（不要加标题、不要加解释）。\n"
        "- 数学公式一律使用 $...$（行内）与 $$...$$（行间）；不要输出 \\( ... \\)、\\[ ... \\] 或 \\begin{equation}...\n"
        "- 如果目标语言是 English/en：输出中不得出现中文（CJK 字符）。\n"
    )

    payload = {"items": items}
    prompt += "\n\nINPUT JSON:\n" + json.dumps(payload, ensure_ascii=False)

    res = safe_chat_completions(
        client=client,
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
        max_tokens=8192,
        timeout_s=timeout_s,
        retries=retries,
        on_attempt=on_attempt,
        response_format={"type": "json_object"},
    )
    if res.error_message:
        raise TranslateCallError(res.error_message, raw=None)

    obj = extract_json_object(res.text)
    if not isinstance(obj, dict) or "items" not in obj or not isinstance(obj.get("items"), list):
        raise TranslateCallError("invalid JSON schema: missing items[]", raw=res.text[:800])

    out_map: Dict[str, List[str]] = {}
    for it in obj["items"]:
        if not isinstance(it, dict):
            continue
        _id = it.get("id")
        segs = it.get("segments")
        if isinstance(_id, str) and isinstance(segs, list) and all(isinstance(s, str) for s in segs):
            out_map[_id] = list(segs)

    # Validate coverage & segment counts
    for it in items:
        _id = it.get("id")
        segs = it.get("segments")
        if not isinstance(_id, str) or not isinstance(segs, list):
            continue
        if _id not in out_map:
            raise TranslateCallError(f"missing translated item: {_id}", raw=res.text[:800])
        if len(out_map[_id]) != len(segs):
            raise TranslateCallError(f"segments length mismatch for {_id}: {len(out_map[_id])} != {len(segs)}", raw=res.text[:800])

    return out_map

# ============================================================
# 2) Ark config / client / retry / timeout
# ============================================================
@dataclass
class ArkResult:
    text: str = ""
    error_message: Optional[str] = None


class TranslateCallError(RuntimeError):
    """Raised when a translate call fails or returns invalid JSON schema."""
    def __init__(self, message: str, raw: Optional[str] = None):
        super().__init__(message)
        self.raw = raw

DEFAULT_ARK_BASE_URL = "https://ark.cn-beijing.volces.com/api/v3"
DEFAULT_ARK_MODEL = "ep-20260203141749-992fx"
DEFAULT_TRANSLATE_MODEL = "ep-20260203211039-n9v2f"  # dedicated translate endpoint
DEFAULT_MML2OMML_XSL_ENV = "MML2OMML_XSL"  # path to MML2OMML.XSL
DEFAULT_OCR_TIMEOUT_S = 120
DEFAULT_TRANSLATE_TIMEOUT_S = 180

def get_api_key() -> str:
    k = None
    try:
        if "ARK_API_KEY" in st.secrets:
            k = st.secrets["ARK_API_KEY"]
    except Exception:
        pass
    k = k or os.environ.get("ARK_API_KEY")
    if not k:
        st.error("缺少 ARK_API_KEY。请在 .streamlit/secrets.toml 或环境变量中配置。")
        st.stop()
    return k

def get_ark_base_url() -> str:
    v = None
    try:
        v = st.secrets.get("ARK_BASE_URL", None)
    except Exception:
        v = None
    return v or os.environ.get("ARK_BASE_URL") or DEFAULT_ARK_BASE_URL

def get_default_model() -> str:
    v = None
    try:
        v = st.secrets.get("ARK_MODEL", None)
    except Exception:
        v = None
    return v or os.environ.get("ARK_MODEL") or os.environ.get("ARK_ENDPOINT_ID") or DEFAULT_ARK_MODEL

def get_timeout_defaults() -> Tuple[int, int]:
    ocr_t = None
    tr_t = None
    try:
        ocr_t = st.secrets.get("ARK_OCR_TIMEOUT_S", None)
        tr_t = st.secrets.get("ARK_TRANSLATE_TIMEOUT_S", None)
    except Exception:
        pass
    ocr_t = ocr_t or os.environ.get("ARK_OCR_TIMEOUT_S") or DEFAULT_OCR_TIMEOUT_S
    tr_t = tr_t or os.environ.get("ARK_TRANSLATE_TIMEOUT_S") or DEFAULT_TRANSLATE_TIMEOUT_S
    return int(ocr_t), int(tr_t)

@st.cache_resource(show_spinner=False)
def get_ark_client_cached(api_key: str, base_url: str, default_timeout_s: int) -> OpenAI:
    return OpenAI(api_key=api_key, base_url=base_url, timeout=default_timeout_s)

def get_ark_client(default_timeout_s: int) -> OpenAI:
    return get_ark_client_cached(get_api_key(), get_ark_base_url(), default_timeout_s)

def _parse_retry_delay_seconds(msg: str) -> Optional[float]:
    m = re.search(r"retry in ([0-9.]+)s", msg, flags=re.IGNORECASE)
    if m:
        return float(m.group(1))
    m = re.search(r'"retryDelay"\s*:\s*"(\d+)s"', msg)
    if m:
        return float(m.group(1))
    return None

def safe_chat_completions(
    client: OpenAI,
    model: str,
    messages: List[dict],
    temperature: float,
    max_tokens: int,
    timeout_s: int,
    retries: int = 6,
    total_timeout_s: Optional[int] = None,
    on_attempt: Optional[Callable[[int, int, str], None]] = None,
    response_format: Optional[dict] = None,
) -> ArkResult:
    """
    - timeout_s: 单次请求超时（传给 OpenAI SDK / httpx）
    - total_timeout_s: 整体“墙钟”超时（包含重试等待）。用于避免卡太久看起来像“死了”。
    - on_attempt: (attempt_idx, retries, phase_msg) 回调，用于在 Streamlit UI 中展示当前在做什么。
    """
    t0 = time.time()
    last_err: Any = None

    for attempt in range(1, retries + 1):
        if total_timeout_s is not None and (time.time() - t0) > float(total_timeout_s):
            return ArkResult(text="", error_message=f"overall timeout after {total_timeout_s}s (last_err={last_err})")

        if on_attempt:
            try:
                on_attempt(attempt, retries, "request")
            except Exception:
                pass

        try:
            # 重要：某些环境里 http 超时/流控可能导致请求“看起来卡死”。
            # 这里用线程包一层硬超时：即使底层库未按预期抛超时，我们也能继续重试并给出可见反馈。
            def _do_req():
                kwargs = dict(
                    model=model,
                    messages=messages,
                    temperature=temperature,
                    max_tokens=max_tokens,
                    timeout=timeout_s,
                )
                if response_format is not None:
                    kwargs["response_format"] = response_format
                try:
                    return client.chat.completions.create(**kwargs)
                except TypeError:
                    # 某些 OpenAI 兼容实现不支持 response_format；回退到普通请求
                    kwargs.pop("response_format", None)
                    return client.chat.completions.create(**kwargs)

            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
                fut = ex.submit(_do_req)
                resp = fut.result(timeout=float(timeout_s) + 5.0)
            txt = ""
            if resp and resp.choices and resp.choices[0].message and resp.choices[0].message.content:
                txt = resp.choices[0].message.content
            return ArkResult(text=(txt or "").strip())

        except concurrent.futures.TimeoutError:
            msg = f"TimeoutError: request exceeded {timeout_s}s"
            last_err = msg
            wait_s = min(2 ** (attempt - 1), 20)
            if on_attempt:
                try:
                    on_attempt(attempt, retries, f"request timeout, sleep {wait_s:.1f}s")
                except Exception:
                    pass
            time.sleep(wait_s)
            continue

        except Exception as e:
            msg = f"{type(e).__name__}: {e}"
            last_err = msg

            # 429 / 限流：按服务端提示或指数退避，但仍受 total_timeout_s 约束
            if ("429" in msg) or ("rate limit" in msg.lower()) or ("RESOURCE_EXHAUSTED" in msg):
                wait_s = _parse_retry_delay_seconds(msg) or min(2 ** (attempt - 1), 60)
                if on_attempt:
                    try:
                        on_attempt(attempt, retries, f"rate-limited, sleep {wait_s:.1f}s")
                    except Exception:
                        pass
                time.sleep(min(wait_s + 0.3, 90.0))
                continue

            # 其它错误：短暂退避
            wait_s = min(2 ** (attempt - 1), 20)
            if on_attempt:
                try:
                    on_attempt(attempt, retries, f"error, sleep {wait_s:.1f}s")
                except Exception:
                    pass
            time.sleep(wait_s)

    return ArkResult(text="", error_message=f"retry exhausted: {last_err}")


# ============================================================
# 3) OCR: image analysis + auto recommend
# ============================================================
def analyze_image_density(img: Image.Image) -> Dict[str, float]:
    g = ImageOps.grayscale(img)
    g_small = g.copy()
    g_small.thumbnail((900, 900))
    w, h = g_small.size

    edges = g_small.filter(ImageFilter.FIND_EDGES)
    eb = edges.point(lambda p: 255 if p > 40 else 0)
    edge_px = sum(1 for p in eb.getdata() if p > 0)
    edge_density = edge_px / float(w * h + 1e-9)

    db = g_small.point(lambda p: 1 if p < 160 else 0)
    dark_px = sum(db.getdata())
    dark_ratio = dark_px / float(w * h + 1e-9)

    return {"edge_density": edge_density, "dark_ratio": dark_ratio, "w": float(img.size[0]), "h": float(img.size[1])}

def recommend_ocr_params(img: Image.Image, mode: str = "Balanced") -> Dict[str, int]:
    m = analyze_image_density(img)
    h = int(m["h"])
    edge = m["edge_density"]
    dark = m["dark_ratio"]
    complexity = 0.6 * edge + 0.4 * dark

    if mode == "Fast":
        max_side, jpeg_q, out_tokens = 1600, 80, 3072
        base_tile, base_overlap = 2000, 120
    elif mode == "Accurate":
        max_side, jpeg_q, out_tokens = 2800, 90, 6144
        base_tile, base_overlap = 1400, 220
    else:
        max_side, jpeg_q, out_tokens = 2200, 85, 4096
        base_tile, base_overlap = 1600, 160

    if complexity > 0.18:
        max_side = min(3200, max_side + 400)
        out_tokens = min(8192, out_tokens + 1024)
        tile_h = max(1000, base_tile - 300)
        overlap = min(320, base_overlap + 80)
        jpeg_q = min(95, jpeg_q + 5)
    elif complexity < 0.10:
        tile_h = min(2400, base_tile + 300)
        overlap = max(80, base_overlap - 40)
    else:
        tile_h = base_tile
        overlap = base_overlap

    if h >= 4000 and complexity > 0.14:
        tile_h = max(900, tile_h - 200)
        out_tokens = min(8192, out_tokens + 512)

    max_side = int(max(800, min(3200, max_side)))
    tile_h = int(max(800, min(2600, tile_h)))
    overlap = int(max(0, min(400, overlap)))
    jpeg_q = int(max(50, min(95, jpeg_q)))
    out_tokens = int(max(1024, min(8192, out_tokens)))

    return {"max_side": max_side, "tile_h": tile_h, "overlap": overlap, "jpeg_q": jpeg_q, "out_tokens": out_tokens}


# ============================================================
# 4) OCR helpers
# ============================================================
def downscale(img: Image.Image, max_side: int) -> Image.Image:
    img = img.convert("RGB")
    w, h = img.size
    s = min(1.0, max_side / float(max(w, h)))
    if s < 1.0:
        img = img.resize((int(w * s), int(h * s)), Image.LANCZOS)
    return img

def pil_to_jpeg_bytes(img: Image.Image, quality: int = 85) -> bytes:
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=quality, optimize=True)
    return buf.getvalue()

def slice_long(img: Image.Image, tile_h: int, overlap: int) -> List[Image.Image]:
    w, h = img.size
    if h <= tile_h:
        return [img]
    overlap = max(0, min(overlap, tile_h // 2))
    step = tile_h - overlap
    out = []
    y = 0
    while y < h:
        y2 = min(y + tile_h, h)
        out.append(img.crop((0, y, w, y2)))
        if y2 >= h:
            break
        y = y + step
    return out

def normalize_md(md: str) -> str:
    md = md.replace("\r\n", "\n").replace("\r", "\n")
    md = re.sub(r"\n{3,}", "\n\n", md).strip()
    return md

def dedupe_tail_head(prev: str, cur: str, max_lines: int = 18) -> str:
    pl = [x.strip() for x in prev.splitlines() if x.strip()]
    cl = [x.strip() for x in cur.splitlines() if x.strip()]
    if not pl or not cl:
        return cur
    k = min(max_lines, len(pl), len(cl))
    for n in range(k, 3, -1):
        if pl[-n:] == cl[:n]:
            raw = cur.splitlines()
            removed = 0
            keep = []
            for line in raw:
                if removed < n and line.strip():
                    removed += 1
                    continue
                keep.append(line)
            return "\n".join(keep).lstrip("\n")
    return cur

def _jpeg_bytes_to_data_url(jpeg_bytes: bytes) -> str:
    b64 = base64.b64encode(jpeg_bytes).decode("utf-8")
    return f"data:image/jpeg;base64,{b64}"

def ocr_image_to_markdown(
    client: OpenAI,
    model: str,
    img: Image.Image,
    max_side: int,
    tile_h: int,
    overlap: int,
    jpeg_q: int,
    out_tokens: int,
    timeout_s: int,
    progress_cb: Optional[Callable[[int, int], None]] = None,
) -> str:
    img = downscale(img, max_side=max_side)
    tiles = slice_long(img, tile_h=tile_h, overlap=overlap)

    chunks: List[str] = []
    total = len(tiles)

    for idx, timg in enumerate(tiles, start=1):
        if progress_cb:
            progress_cb(idx, total)

        b = pil_to_jpeg_bytes(timg, quality=jpeg_q)
        data_url = _jpeg_bytes_to_data_url(b)

        messages = [{
            "role": "user",
            "content": [
                {"type": "text", "text": OCR_PROMPT_ZH},
                {"type": "image_url", "image_url": {"url": data_url}},
            ],
        }]

        res = safe_chat_completions(
            client=client,
            model=model,
            messages=messages,
            temperature=0.2,
            max_tokens=out_tokens,
            timeout_s=timeout_s,
            retries=6,
        )
        if res.error_message:
            st.error(f"OCR 调用失败：\n\n{res.error_message}")
            st.stop()

        md = normalize_md(res.text)
        if chunks:
            md = dedupe_tail_head(chunks[-1], md)
        chunks.append(md)

    return normalize_md("\n\n".join([c for c in chunks if c.strip()]))


# ============================================================
# 4.5) PDF -> images (PyMuPDF)
# ============================================================
def parse_page_range(spec: str, *, max_pages: Optional[int] = None) -> Optional[List[int]]:
    """Parse human page range like '1-3,5,8-10' into 0-based page indices.
    Returns None if spec is empty/invalid (meaning: use default behavior).
    """
    if spec is None:
        return None
    s = str(spec).strip()
    if not s:
        return None
    s = s.replace("，", ",").replace(" ", "")
    out: List[int] = []
    for part in s.split(","):
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            if not a.isdigit() or not b.isdigit():
                return None
            lo, hi = int(a), int(b)
            if lo <= 0 or hi <= 0:
                return None
            if lo > hi:
                lo, hi = hi, lo
            for p in range(lo, hi + 1):
                out.append(p - 1)
        else:
            if not part.isdigit():
                return None
            p = int(part)
            if p <= 0:
                return None
            out.append(p - 1)
    # de-dup and clamp
    out = sorted(set(out))
    if max_pages is not None:
        out = [p for p in out if 0 <= p < max_pages]
    return out if out else None


# ============================================================
# 4.5) PDF -> images (PyMuPDF)
# ============================================================
def pdf_bytes_to_images(
    pdf_bytes: bytes,
    dpi: int = 300,
    max_pages: Optional[int] = None,
    page_indices: Optional[List[int]] = None,
) -> List[Image.Image]:
    """Render PDF pages to PIL Images.
    - If page_indices is provided (0-based), only render those pages.
    - Else render first N pages (N=max_pages or all).
    """
    if not HAVE_PYMUPDF:
        raise RuntimeError("缺少依赖 pymupdf。请先 pip install pymupdf")
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    imgs: List[Image.Image] = []
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)

    if page_indices:
        indices = [i for i in page_indices if 0 <= i < doc.page_count]
    else:
        n = doc.page_count if max_pages is None else min(doc.page_count, max_pages)
        indices = list(range(n))

    for i in indices:
        page = doc.load_page(i)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
        imgs.append(img)
    return imgs


# ============================================================
# 5) LaTeX code style for OCR result
# ============================================================
DISPLAY_MATH_RE = re.compile(r"(?s)\$\$(.+?)\$\$")
INLINE_MATH_RE = re.compile(r"(?<!\$)\$([^$\n]+)\$(?!\$)")

def display_to_equation_env(md: str) -> str:
    def repl(m):
        body = m.group(1).strip()
        return "\\begin{equation}\n" + body + "\n\\end{equation}"
    return DISPLAY_MATH_RE.sub(repl, md)

def inline_to_paren(md: str) -> str:
    def repl(m):
        body = m.group(1).strip()
        return "\\(" + body + "\\)"
    return INLINE_MATH_RE.sub(repl, md)

def md_to_latex_code_style(md: str) -> str:
    md2 = normalize_md(md)
    md2 = display_to_equation_env(md2)
    md2 = inline_to_paren(md2)
    eq_env = re.compile(r"(?s)(\\begin\{equation\}.*?\\end\{equation\})")
    md2 = eq_env.sub(lambda m: "\n```latex\n" + m.group(1).strip() + "\n```\n", md2)
    return normalize_md(md2)


# ============================================================
# 6) DOCX translate (preserve layout, best-effort) + equation replacement
# ============================================================
MATH_TOKEN_RE = re.compile(
    r"(\$\$.*?\$\$|\$[^$\n]+\$|\\begin\{equation\}.*?\\end\{equation\}|\\\(.+?\\\))",
    re.DOTALL
)

def protect_math(text: str) -> Tuple[str, Dict[str, str]]:
    mapping: Dict[str, str] = {}
    k = 0
    def repl(m):
        nonlocal k
        token = f"__MATH_{k}__"
        mapping[token] = m.group(0)
        k += 1
        return token
    out = MATH_TOKEN_RE.sub(repl, text)
    return out, mapping

def restore_tokens(text: str, mapping: Dict[str, str]) -> str:
    for tok, val in mapping.items():
        text = text.replace(tok, val)
    return text

def iter_table_paragraphs(table: Table) -> List[Paragraph]:
    out: List[Paragraph] = []
    for row in table.rows:
        for cell in row.cells:
            out.extend(cell.paragraphs)
            for t2 in cell.tables:
                out.extend(iter_table_paragraphs(t2))
    return out

def iter_part_paragraphs(part) -> List[Paragraph]:
    """part: doc, header, footer"""
    ps: List[Paragraph] = []
    try:
        ps.extend(part.paragraphs)
    except Exception:
        pass
    try:
        for table in part.tables:
            ps.extend(iter_table_paragraphs(table))
    except Exception:
        pass
    return ps

def iter_all_paragraphs_extended(doc: Document) -> List[Paragraph]:
    """
    ✅ 覆盖：正文 paragraphs + tables + 每个 section 的 header/footer（含表格）
    仍不覆盖：textbox/shape、批注、脚注尾注（python-docx 限制）
    """
    ps: List[Paragraph] = []
    ps.extend(iter_part_paragraphs(doc))

    # headers/footers
    try:
        for sec in doc.sections:
            ps.extend(iter_part_paragraphs(sec.header))
            ps.extend(iter_part_paragraphs(sec.footer))
            # first/even page headers/footers（若存在）
            if hasattr(sec, "first_page_header"):
                ps.extend(iter_part_paragraphs(sec.first_page_header))
            if hasattr(sec, "first_page_footer"):
                ps.extend(iter_part_paragraphs(sec.first_page_footer))
            if hasattr(sec, "even_page_header"):
                ps.extend(iter_part_paragraphs(sec.even_page_header))
            if hasattr(sec, "even_page_footer"):
                ps.extend(iter_part_paragraphs(sec.even_page_footer))
    except Exception:
        pass
    return ps

def extract_math_from_docx_with_pandoc(docx_bytes: bytes) -> List[Tuple[str, str]]:
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        f = td / "in.docx"
        f.write_bytes(docx_bytes)
        md = pypandoc.convert_file(str(f), to="markdown", format="docx", extra_args=["--wrap=none"])
        md = md.replace("\r\n", "\n").replace("\r", "\n")

    matches: List[Tuple[int, int, str, str]] = []
    for m in re.finditer(r"(?s)\$\$(.+?)\$\$", md):
        matches.append((m.start(), m.end(), "display", m.group(1).strip()))
    for m in re.finditer(r"(?<!\$)\$([^$\n]+)\$(?!\$)", md):
        matches.append((m.start(), m.end(), "inline", m.group(1).strip()))
    matches.sort(key=lambda x: x[0])
    return [(k, b) for _, _, k, b in matches]

def insert_text_run_at_paragraph_child(p_elm, idx: int, text: str):
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p_elm.insert(idx, r)

def replace_omml_with_latex_code(doc: Document, math_seq: List[Tuple[str, str]], use_equation_env: bool = True) -> int:
    all_ps = iter_all_paragraphs_extended(doc)
    seq_idx = 0
    replaced = 0

    for p in all_ps:
        p_elm = p._p
        nodes = p_elm.xpath(".//*[local-name()='oMath' or local-name()='oMathPara']")
        for node in nodes:
            if seq_idx >= len(math_seq):
                return replaced
            kind, body = math_seq[seq_idx]
            seq_idx += 1

            if use_equation_env:
                latex_text = (f"\\begin{{equation}} {body} \\end{{equation}}") if kind == "display" else (f"\\({body}\\)")
            else:
                latex_text = (f"$$\n{body}\n$$") if kind == "display" else (f"${body}$")

            parent = node.getparent()
            if parent is None:
                continue

            try:
                idx_in_parent = list(parent).index(node)
            except Exception:
                idx_in_parent = None

            if parent is not p_elm or idx_in_parent is None:
                insert_text_run_at_paragraph_child(p_elm, len(p_elm), latex_text)
            else:
                insert_text_run_at_paragraph_child(parent, idx_in_parent, latex_text)

            try:
                parent.remove(node)
            except Exception:
                pass

            replaced += 1

    return replaced


# ---------------------------
# LaTeX ($...$ / $$...$$ / \(..\) / \[..]) -> Word equation (OMML)
# ---------------------------
LATEX_INLINE_RE = re.compile(
    r"(\$\$.*?\$\$|\$[^$\n]+\$|\\\(.+?\\\)|\\\[.+?\\\])",
    re.DOTALL,
)

def _strip_latex_delims(s: str) -> Tuple[str, bool]:
    s = s.strip()
    if s.startswith("$$") and s.endswith("$$"):
        return s[2:-2].strip(), True
    if s.startswith("$") and s.endswith("$"):
        return s[1:-1].strip(), False
    if s.startswith(r"\(") and s.endswith(r"\)"):
        return s[2:-2].strip(), False
    if s.startswith(r"\[") and s.endswith(r"\]"):
        return s[2:-2].strip(), True
    return s, False

def _guess_office_mml2omml_paths() -> List[str]:
    # Common installs; user can override via sidebar/env.
    return [
        r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\Office16\MML2OMML.XSL",
    ]

@lru_cache(maxsize=1)
def _load_mml2omml_xslt(xsl_path: str):
    if not HAVE_LATEX_OMML:
        raise RuntimeError("Missing deps: lxml / latex2mathml")
    if not xsl_path:
        # try guess
        for p in _guess_office_mml2omml_paths():
            if os.path.exists(p):
                xsl_path = p
                break
    if not xsl_path or not os.path.exists(xsl_path):
        raise RuntimeError(
            "MML2OMML.XSL not found. Provide its path in sidebar (or env MML2OMML_XSL)."
        )
    xslt_doc = etree.parse(xsl_path)
    return etree.XSLT(xslt_doc)

def latex_to_omml_element(latex_expr: str, xsl_path: str):
    """Return an lxml element (OMML) for insertion into docx XML."""
    if not HAVE_LATEX_OMML:
        raise RuntimeError("This feature requires: pip install lxml latex2mathml")
    body, is_block = _strip_latex_delims(latex_expr)
    # latex2mathml returns <math xmlns="http://www.w3.org/1998/Math/MathML">...</math>
    mml = latex2mathml.converter.convert(body)
    # Ensure no newlines (Word is sensitive per some implementations)
    mml = re.sub(r"\s+", " ", mml).strip()
    mml_root = etree.fromstring(mml.encode("utf-8"))
    xslt = _load_mml2omml_xslt(xsl_path)
    omml_tree = xslt(mml_root)
    omml_root = omml_tree.getroot()
    # Word accepts either m:oMath or m:oMathPara; keep as produced by stylesheet.
    return omml_root, is_block

def _clone_rPr(src_r):
    # clone run properties to keep styling for newly inserted text runs
    rpr = src_r.find(qn("w:rPr"))
    if rpr is None:
        return None
    return deepcopy(rpr)

def convert_inline_latex_to_omml_in_doc(doc: Document, xsl_path: str, log: Optional[Callable[[str], None]] = None) -> int:
    """Best-effort: convert LaTeX math delimited with $...$ etc into OMML equations inside a docx.

    Limitations:
    - Only detects math fully contained inside a single run.
    - Keeps paragraph structure; tries to preserve run formatting for surrounding text.
    """
    if not HAVE_LATEX_OMML:
        raise RuntimeError("Missing deps: lxml / latex2mathml")
    converted = 0
    for p in doc.paragraphs:
        # iterate by underlying XML runs to allow insertion
        r_elems = list(p._p.findall(qn("w:r")))
        for r in r_elems:
            t = r.find(qn("w:t"))
            if t is None or not t.text:
                continue
            text = t.text
            # find first match (one per run; iterate after mutation)
            m = LATEX_INLINE_RE.search(text)
            if not m:
                continue
            before = text[:m.start()]
            expr = m.group(0)
            after = text[m.end():]

            try:
                omml_elem, _is_block = latex_to_omml_element(expr, xsl_path)
            except Exception as e:
                if log:
                    log(f"公式转换失败（跳过）：{type(e).__name__}: {e} | expr={expr[:80]}")
                continue

            # Replace current run text with "before"
            t.text = before

            # Insert OMML element right after this run
            # Note: OMML uses its own namespace (m:). etree element is fine to append.
            r.addnext(omml_elem)

            # Insert "after" as a new run, trying to keep same rPr
            if after:
                new_r = OxmlElement("w:r")
                rpr_clone = _clone_rPr(r)
                if rpr_clone is not None:
                    new_r.append(rpr_clone)
                new_t = OxmlElement("w:t")
                new_t.text = after
                new_r.append(new_t)
                omml_elem.addnext(new_r)

            converted += 1
    return converted


def chunk_items_for_api(items: List[dict], max_chars: int = 12000) -> List[List[dict]]:
    batches: List[List[dict]] = []
    cur: List[dict] = []
    cur_len = 0
    for it in items:
        s = json.dumps(it, ensure_ascii=False)
        if cur and cur_len + len(s) > max_chars:
            batches.append(cur)
            cur = []
            cur_len = 0
        cur.append(it)
        cur_len += len(s)
    if cur:
        batches.append(cur)
    return batches

def extract_json_object(text: str) -> dict:
    if not text:
        raise ValueError("empty response")
    s = text.strip()

    fence = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", s, flags=re.DOTALL | re.IGNORECASE)
    if fence:
        s = fence.group(1).strip()

    try:
        return json.loads(s)
    except Exception:
        pass

    start = s.find("{")
    if start == -1:
        raise ValueError("no '{' found in response")

    depth = 0
    in_str = False
    esc = False
    for i in range(start, len(s)):
        ch = s[i]
        if in_str:
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == '"':
                in_str = False
        else:
            if ch == '"':
                in_str = True
            elif ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    candidate = s[start:i + 1]
                    return json.loads(candidate)

    raise ValueError("no complete JSON object found (unbalanced braces)")

def _ark_extract_output_text(resp_obj: dict) -> str:
    """Extract assistant text from an Ark/OpenAI-style Responses API payload.

    Ark /api/v3/responses follows the OpenAI Responses schema:
      resp.output -> list[message] -> content -> list[{type: "output_text", text: "..."}]
    Some gateways may also return ChatCompletions-style `choices`.
    """
    texts: List[str] = []
    if isinstance(resp_obj, dict):
        out = resp_obj.get("output")
        if isinstance(out, list):
            for msg in out:
                if not isinstance(msg, dict):
                    continue
                content = msg.get("content")
                if not isinstance(content, list):
                    continue
                for part in content:
                    if isinstance(part, dict) and part.get("type") == "output_text" and isinstance(part.get("text"), str):
                        texts.append(part["text"])

        if not texts:
            choices = resp_obj.get("choices")
            if isinstance(choices, list):
                for ch in choices:
                    try:
                        t = ch["message"]["content"]
                        if isinstance(t, str):
                            texts.append(t)
                    except Exception:
                        pass

        if not texts and isinstance(resp_obj.get("text"), str):
            texts.append(resp_obj["text"])
    return "".join(texts).strip()


def ark_translate_text_via_responses(
    api_key: str,
    base_url: str,
    text: str,
    src_lang: str,
    dst_lang: str,
    timeout_s: int,
    model: str = DEFAULT_TRANSLATE_MODEL,
    retries: int = 6,
    on_attempt: Optional[Callable[[int, int, str], None]] = None,
) -> str:
    """Call Ark Responses API translation model using translation_options.

    Docs: /api/v3/responses supports input_text.translation_options.source_language/target_language.
    Response: output.content.type=output_text, output.content.text holds the text.
    """
    url = base_url.rstrip("/") + "/responses"

    # Ark translation expects language codes; we accept friendly labels from UI and map to codes.
    lang_map = {
        "Auto": None,
        "Chinese": "zh",
        "English": "en",
        "Japanese": "ja",
        "Korean": "ko",
        "Spanish": "es",
        "French": "fr",
        "German": "de",
        "Portuguese": "pt",
        "Russian": "ru",
        "Thai": "th",
        "Vietnamese": "vi",
        "Arabic": "ar",
        "Czech": "cs",
        "Danish": "da",
    }

    src_code = lang_map.get(src_lang, src_lang)
    dst_code = lang_map.get(dst_lang, dst_lang) or dst_lang
    if dst_code in (None, "", "Auto"):
        raise TranslateCallError("target language must be specified (not Auto)")

    trans_opts: Dict[str, str] = {"target_language": dst_code}
    if src_code:
        trans_opts["source_language"] = src_code

    payload = {
        "model": model,
        "input": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "input_text",
                        "text": text,
                        "translation_options": trans_opts,
                    }
                ],
            }
        ],
    }

    last_err: Optional[str] = None
    for attempt in range(1, retries + 1):
        if on_attempt:
            on_attempt(attempt, retries, "send")

        try:
            with httpx.Client(timeout=timeout_s) as hc:
                r = hc.post(
                    url,
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json",
                    },
                    json=payload,
                )

            if r.status_code >= 400:
                try:
                    err_obj = r.json()
                except Exception:
                    err_obj = None
                msg = None
                if isinstance(err_obj, dict):
                    err = err_obj.get("error")
                    if isinstance(err, dict):
                        msg = err.get("message") or err.get("msg")
                    msg = msg or err_obj.get("message") or err_obj.get("msg")
                msg = msg or f"HTTP {r.status_code}: {r.text[:400]}"
                last_err = msg

                if r.status_code in (408, 409, 429) or (500 <= r.status_code < 600):
                    delay = _parse_retry_delay_seconds(msg) or min(8.0, 0.8 * (2 ** (attempt - 1)))
                    if on_attempt:
                        on_attempt(attempt, retries, f"backoff {delay:.1f}s")
                    time.sleep(delay)
                    continue
                raise TranslateCallError(f"translate http error: {msg}")

            obj = r.json()
            out_text = _ark_extract_output_text(obj)
            if not out_text:
                raise TranslateCallError("empty translate response text")
            return out_text

        except (httpx.TimeoutException, httpx.TransportError) as e:
            last_err = f"{type(e).__name__}: {e}"
            delay = min(8.0, 0.8 * (2 ** (attempt - 1)))
            if on_attempt:
                on_attempt(attempt, retries, f"network backoff {delay:.1f}s")
            time.sleep(delay)

    raise TranslateCallError(f"translate failed after retries: {last_err}")


def doubao_translate_items(
    client: OpenAI,
    model: str,
    items: List[dict],
    src_lang: str,
    dst_lang: str,
    timeout_s: int,
    on_attempt: Optional[Callable[[int, int, str], None]] = None,
    debug_sink: Optional[Callable[[str], None]] = None,
) -> Dict[str, List[str]]:
    """Translate a batch of items via the dedicated translation model (Responses API).

    IMPORTANT (per requirement):
    - When translation is enabled, we ALWAYS use DEFAULT_TRANSLATE_MODEL = ep-20260203211039-n9v2f.
    - We do NOT use other models for translation and do NOT ask the model to output JSON.
    - We translate each segment directly so $...$/__MATH_x__ placeholders are preserved by protect_math().
    """
    api_key = get_api_key()
    base_url = get_ark_base_url()
    translate_model = DEFAULT_TRANSLATE_MODEL

    # Flatten tasks
    tasks: List[Tuple[str, int, str]] = []
    out_buf: Dict[str, List[Optional[str]]] = {}

    for it in items:
        _id = it.get("id")
        segs = it.get("segments")
        if not isinstance(_id, str) or not isinstance(segs, list):
            continue
        out_buf[_id] = [None] * len(segs)
        for j, seg in enumerate(segs):
            tasks.append((_id, j, seg if isinstance(seg, str) else ""))

    # Concurrency: keep conservative to avoid rate-limit bursts.
    max_workers = min(4, max(2, (os.cpu_count() or 4) // 2))
    errors: List[str] = []

    def _do_one(t: Tuple[str, int, str]) -> Tuple[str, int, str]:
        _id, j, seg = t
        if not seg.strip():
            return _id, j, seg

        translated = ark_translate_text_via_responses(
            api_key=api_key,
            base_url=base_url,
            text=seg,
            src_lang=src_lang,
            dst_lang=dst_lang,
            timeout_s=timeout_s,
            model=translate_model,
            retries=6,
            on_attempt=on_attempt,
        )
        if debug_sink:
            debug_sink(translated[:400])
        return _id, j, translated

    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as ex:
        futs = [ex.submit(_do_one, t) for t in tasks]
        for fut in concurrent.futures.as_completed(futs):
            try:
                _id, j, translated = fut.result()
                if _id in out_buf and 0 <= j < len(out_buf[_id]):
                    out_buf[_id][j] = translated
            except Exception as e:
                errors.append(f"{type(e).__name__}: {e}")

    # Retry failed segments sequentially once (often fixes transient 429).
    if any(v is None for segs in out_buf.values() for v in segs):
        for _id, segs in list(out_buf.items()):
            for j, val in enumerate(segs):
                if val is not None:
                    continue
                seg = next((s for (iid, jj, s) in tasks if iid == _id and jj == j), "")
                try:
                    out_buf[_id][j] = ark_translate_text_via_responses(
                        api_key=api_key,
                        base_url=base_url,
                        text=seg,
                        src_lang=src_lang,
                        dst_lang=dst_lang,
                        timeout_s=timeout_s,
                        model=translate_model,
                        retries=6,
                        on_attempt=on_attempt,
                    )
                except Exception as e:
                    errors.append(f"{type(e).__name__}: {e}")

    if errors and any(v is None for segs in out_buf.values() for v in segs):
        raise TranslateCallError("; ".join(errors[:3]) + (" ..." if len(errors) > 3 else ""))

    final: Dict[str, List[str]] = {}
    for _id, segs in out_buf.items():
        if any(s is None for s in segs):
            raise TranslateCallError(f"missing translated segments for {_id}")
        final[_id] = [s or "" for s in segs]
    return final

def estimate_auto_batch_chars(items: List[dict], target_batches: int, clamp_min: int = 4000, clamp_max: int = 20000) -> int:
    """
    自动分批大小：
    - 先估 total_json_chars
    - 再除以 target_batches 得到 max_batch_chars
    - clamp 到 [clamp_min, clamp_max]
    """
    if not items:
        return 12000
    total = 0
    for it in items:
        total += len(json.dumps(it, ensure_ascii=False))
    est = int(total / max(1, target_batches))
    return int(max(clamp_min, min(clamp_max, est)))


def translate_items_adaptive(
    client: OpenAI,
    model: str,
    items: List[dict],
    src_lang: str,
    dst_lang: str,
    timeout_s: int,
    on_attempt: Optional[Callable[[int, int, str], None]] = None,
    debug_sink: Optional[Callable[[str], None]] = None,
    max_depth: int = 10,
) -> Dict[str, List[str]]:
    """Translate with automatic batch splitting on timeout/invalid JSON.

    This is the main fix for “small text works, large text times out / hangs”.
    Strategy:
    - Try translating the batch as-is.
    - If it fails (timeout, rate-limit, invalid JSON), split the batch into halves and retry.
    - Merge results. Stops splitting when batch has 1 item or max_depth reached.
    """
    if not items:
        return {}
    try:
        return doubao_translate_items(
            client=client,
            model=model,
            items=items,
            src_lang=src_lang,
            dst_lang=dst_lang,
            timeout_s=timeout_s,
            on_attempt=on_attempt,
            debug_sink=debug_sink,
        )
    except TranslateCallError as e:
        # Only split if we can
        if len(items) <= 1 or max_depth <= 0:
            raise
        mid = len(items) // 2
        left = translate_items_adaptive(client, model, items[:mid], src_lang, dst_lang, timeout_s,
                                       on_attempt=on_attempt, debug_sink=debug_sink, max_depth=max_depth - 1)
        right = translate_items_adaptive(client, model, items[mid:], src_lang, dst_lang, timeout_s,
                                        on_attempt=on_attempt, debug_sink=debug_sink, max_depth=max_depth - 1)
        left.update(right)
        return left

def translate_docx_in_place(
    doc: Document,
    client: OpenAI,
    model: str,
    src_lang: str,
    dst_lang: str,
    max_batch_chars: int,
    timeout_s: int,
    progress_cb: Optional[Callable[[int, int], None]] = None,
    diagnostic_cb: Optional[Callable[[Dict[str, Any]], None]] = None,
    attempt_msg_cb: Optional[Callable[[str], None]] = None,
):
    all_ps = iter_all_paragraphs_extended(doc)

    items: List[dict] = []
    para_refs: Dict[str, Tuple[List[Any], List[Dict[str, str]]]] = {}
    pid = 0

    total_runs = 0
    total_visible_chars = 0

    for p in all_ps:
        runs = [r for r in p.runs if r.text is not None and r.text != ""]
        if not runs:
            continue

        segs = []
        maps = []
        for r in runs:
            total_runs += 1
            total_visible_chars += len(r.text or "")
            protected, mp = protect_math(r.text)
            segs.append(protected)
            maps.append(mp)

        if not "".join(segs).strip():
            continue

        pid += 1
        item_id = f"p{pid}"
        items.append({"id": item_id, "segments": segs})
        para_refs[item_id] = (runs, maps)

    if diagnostic_cb is not None:
        diagnostic_cb({
            "paragraphs_scanned": len(all_ps),
            "items_built": len(items),
            "runs_counted": total_runs,
            "visible_chars_counted": total_visible_chars,
        })

    if not items:
        # 关键：明确提示为什么“没有翻译”
        st.warning(
            "没有抓到可翻译的正文文本（items=0）。\n\n"
            "常见原因：\n"
            "1) 内容在文本框/形状（python-docx 默认无法读取）；\n"
            "2) 内容主要在批注/脚注/尾注；\n"
            "3) 内容是嵌入对象或图片。\n\n"
            "建议：\n"
            "- 若是文本框：优先用 Tab③ Pandoc 导出，或将文本框内容复制到正文段落后再翻译；\n"
            "- 或改走 PDF/图片 OCR 路线。"
        )
        return

    batches = chunk_items_for_api(items, max_chars=max_batch_chars)
    total = len(batches)

    for bi, batch in enumerate(batches, start=1):
        if progress_cb:
            progress_cb(bi, total)

        if attempt_msg_cb:
            try:
                approx_chars = sum(len(json.dumps(it, ensure_ascii=False)) for it in batch)
            except Exception:
                approx_chars = -1
            attempt_msg_cb(
                f"批次 {bi}/{total}：发送翻译请求（items={len(batch)}，approx_chars={approx_chars}）"
            )

        translated_map = translate_items_adaptive(
            client=client,
            model=model,
            items=batch,
            src_lang=src_lang,
            dst_lang=dst_lang,
            timeout_s=timeout_s,
            on_attempt=(lambda a, n, ph: attempt_msg_cb(f"批次 {bi}/{total}：请求尝试 {a}/{n} · {ph}") if attempt_msg_cb else None),
            debug_sink=(lambda raw: attempt_msg_cb(f"批次 {bi}/{total}：已收到原始输出（len={len(raw)})") if attempt_msg_cb else None),
        )

        if attempt_msg_cb:
            attempt_msg_cb(f"批次 {bi}/{total}：收到响应，开始解析/写回…")

        for it in batch:
            item_id = it["id"]
            if item_id not in translated_map:
                continue

            runs, maps = para_refs[item_id]
            out_segs = translated_map[item_id]

            if len(out_segs) != len(runs):
                whole = " ".join(out_segs)
                whole = restore_tokens(whole, {k: v for mp in maps for k, v in mp.items()})
                runs[0].text = whole
                for r in runs[1:]:
                    r.text = ""
                continue

            for r, seg, mp in zip(runs, out_segs, maps):
                r.text = restore_tokens(seg, mp)

        if attempt_msg_cb:
            attempt_msg_cb(f"批次 {bi}/{total}：写回完成")


# ============================================================
# 7) Export helpers
# ============================================================
def doc_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def pandoc_md_to_docx(md: str) -> bytes:
    md = normalize_md(md) + "\n"
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "out.docx"
        fmt = "markdown+fenced_code_blocks+tex_math_dollars"
        pypandoc.convert_text(md, to="docx", format=fmt, outputfile=str(out))
        return out.read_bytes()


# ============================================================
# ============================================================
# 7.5) Pandoc math normalization helpers (for thesis-friendly editing)
# ============================================================
PANDOC_DISPLAY_MATH_BRACKET_RE = re.compile(r"(?s)\\\[(.+?)\\\]")  # \[ ... \]
PANDOC_DISPLAY_MATH_DOLLAR_RE = re.compile(r"(?s)\$\$(.+?)\$\$")    # $$ ... $$
PANDOC_INLINE_MATH_PAREN_RE = re.compile(r"(?s)\\\((.+?)\\\)")    # \( ... \)
PANDOC_INLINE_MATH_DOLLAR_RE = re.compile(r"(?s)(?<!\$)\$([^$\n]+)\$(?!\$)")  # $ ... $

def normalize_pandoc_math(
    text: str,
    *,
    display_style: str = "equation",   # equation | equation* | bracket | dollars
    inline_style: str = "paren",       # paren | dollars
) -> str:
    """Normalize math delimiters produced by pandoc for easier paper editing.

    - Pandoc LaTeX writer commonly emits inline math as \(..\) and display math as \[..\].
    - Some sources may already contain $$...$$. We normalize both.
    """
    if not text:
        return text

    t = text.replace("\r\n", "\n").replace("\r", "\n")

    def _to_equation(body: str, starred: bool) -> str:
        env = "equation*" if starred else "equation"
        body = body.strip()
        return f"\\begin{{{env}}}\n{body}\n\\end{{{env}}}"

    def _display_repl(m):
        body = (m.group(1) or "").strip()
        if display_style == "equation":
            return _to_equation(body, starred=False)
        if display_style == "equation*":
            return _to_equation(body, starred=True)
        if display_style == "dollars":
            return f"$$\n{body}\n$$"
        return f"\\[\n{body}\n\\]"

    # Normalize display math first (both \[\] and $$ $$)
    t = PANDOC_DISPLAY_MATH_BRACKET_RE.sub(_display_repl, t)
    t = PANDOC_DISPLAY_MATH_DOLLAR_RE.sub(_display_repl, t)

    def _inline_repl(m):
        body = (m.group(1) or "").strip()
        if inline_style == "dollars":
            return f"${body}$"
        return f"\\({body}\\)"

    # Normalize inline math (both \(\) and $ $)
    t = PANDOC_INLINE_MATH_PAREN_RE.sub(_inline_repl, t)
    t = PANDOC_INLINE_MATH_DOLLAR_RE.sub(lambda m: _inline_repl(m), t)
    return t


# ---- Optional: improve Word OMML visual fidelity for "text-like" subscripts/superscripts ----
# Example: P_seed  -> P_{\mathrm{seed}}
MATH_DOLLAR_BLOCK_RE = re.compile(r"(?s)\$\$(.+?)\$\$")
MATH_DOLLAR_INLINE_RE = re.compile(r"(?s)(?<!\$)\$([^$\n]+)\$(?!\$)")
SUBSUP_TEXT_RE = re.compile(r"([_^])([A-Za-z]{2,})\b")

def _fix_subsup_text_style_in_math(expr: str) -> str:
    def repl(m):
        op = m.group(1)  # _ or ^
        word = m.group(2)
        return f"{op}{{\\mathrm{{{word}}}}}"
    return SUBSUP_TEXT_RE.sub(repl, expr)

def normalize_math_text_style(md: str) -> str:
    """Heuristic: inside $...$/$$...$$ convert bare multi-letter subscripts/superscripts to \mathrm{...}."""
    if not md:
        return md

    def fix_block(m):
        body2 = _fix_subsup_text_style_in_math(m.group(1))
        return "$$\n" + body2.strip() + "\n$$"

    def fix_inline(m):
        body2 = _fix_subsup_text_style_in_math(m.group(1))
        return "$" + body2.strip() + "$"

    md2 = MATH_DOLLAR_BLOCK_RE.sub(fix_block, md)
    md2 = MATH_DOLLAR_INLINE_RE.sub(fix_inline, md2)
    return md2




def _sanitize_tex_math_for_pandoc(md: str) -> str:
    """Make TeX math in Markdown more pandoc-friendly before markdown->docx.

    Why this exists:
    - Pandoc's TeX math with $$...$$ does *not* allow blank lines inside the block,
      otherwise it may terminate display math early and the remainder becomes normal text. 
    - When coming from DOCX, pandoc's markdown writer can escape backslashes as
      \\alpha (meaning \alpha). Inside math, we want to unescape command starters
      (\\alpha -> \alpha) but keep real LaTeX linebreaks (\\) intact.

    We only touch content *inside* $...$ / $$...$$.
    """
    if not md:
        return md

    def _fix_math_body(body: str) -> str:
        # 1) remove blank lines inside display math
        body = re.sub(r"\n[ \t]*\n+", "\n", body)

        # 2) unescape backslashes only when it's likely a command starter:
        #    \\alpha, \\rho, \\_, \\{  -> \alpha, \rho, \_, \{
        #    but keep \\ (linebreak) when followed by whitespace/newline.
        body = re.sub(r"\\\\(?=[A-Za-z_{])", r"\\", body)
        return body

    def _disp(m):
        body = _fix_math_body(m.group(1)).strip(" \n\t")
        return "$$\n" + body + "\n$$"

    def _inline(m):
        body = _fix_math_body(m.group(1)).replace("\n", " ").strip()
        return "$" + body + "$"


    md2 = MATH_DOLLAR_BLOCK_RE.sub(_disp, md)
    md2 = MATH_DOLLAR_INLINE_RE.sub(_inline, md2)
    return md2
# ============================================================
# 7.6) DOCX round-trip: $...$ / $$...$$ text math -> editable Word equations (OMML)
#      (Pandoc parses TeX math and writes OMML)
# ============================================================

# $$ ... $$ and $ ... $ (inline) in markdown
MATH_DOLLAR_BLOCK_RE = re.compile(r"(?s)\$\$(.+?)\$\$")
MATH_DOLLAR_INLINE_RE = re.compile(r"(?s)(?<!\$)\$([^$\n]+)\$(?!\$)")

# Heuristic: sub/sup is plain text token (letters length>=2) -> wrap with \mathrm{...}
# Example: P_c(t_k^+) -> c and k are single-letter (keep), but "seed" or "sc" should become \mathrm{seed}
SUBSUP_TEXT_RE = re.compile(r"([_^])([A-Za-z]{2,})\b")

def _fix_subsup_text_style_in_math(expr: str) -> str:
    def repl(m):
        op = m.group(1)
        word = m.group(2)
        return f"{op}{{\\mathrm{{{word}}}}}"
    return SUBSUP_TEXT_RE.sub(repl, expr)

def normalize_math_text_style(md: str) -> str:
    """Fix a common Pandoc/Word math nuance:
    In TeX, multi-letter tokens in sub/sup are treated as variables (italic) unless wrapped.
    For academic docs, users often intend them as text (e.g., P_seed, rho_h, t_k).
    We heuristically wrap multi-letter sub/sup tokens with \mathrm{...}.
    """
    if not md:
        return md

    def fix_block(m):
        body = m.group(1)
        body2 = _fix_subsup_text_style_in_math(body)
        return "$$\n" + body2.strip() + "\n$$"

    def fix_inline(m):
        body = m.group(1)
        body2 = _fix_subsup_text_style_in_math(body)
        return "$" + body2.strip() + "$"

    t = md
    t = MATH_DOLLAR_BLOCK_RE.sub(fix_block, t)
    t = MATH_DOLLAR_INLINE_RE.sub(fix_inline, t)
    return t



def docx_roundtrip_make_equations_editable(docx_bytes: bytes) -> bytes:
    """Round-trip a DOCX through Pandoc Markdown to convert $...$/$$...$$ into native Word equations (OMML).

    Pipeline:
      1) docx -> markdown+tex_math_dollars (keeps math in plain text)
      2) normalize escapes + math delimiters
      3) markdown+tex_math_dollars -> docx (Pandoc emits OMML for TeX math)
      4) use --reference-doc to preserve the original styles as much as possible
    """
    if not docx_bytes:
        return docx_bytes

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        in_path = td / "in.docx"
        in_path.write_bytes(docx_bytes)

        # 1) docx -> markdown (keep $...$/$$...$$ as TeX math)
        md = pypandoc.convert_file(
            str(in_path),
            to="markdown+tex_math_dollars",
            format="docx",
            extra_args=["--wrap=none"],
        )
        md = md.replace("\r\n", "\n").replace("\r", "\n")

        # 2) Undo pandoc escaping so TeX math stays parseable for the next step.
        #    (In academic docs, literal dollar signs are rare; this is a pragmatic choice.)
        md = md.replace(r"\$", "$")
        md = md.replace("\\\\", "\\")  # \\ -> \

        # 3) Normalize delimiters and text-style in subscripts/superscripts
        md = normalize_pandoc_math(md, display_style="dollars", inline_style="dollars")
        md = normalize_math_text_style(md)

        # 4) markdown -> docx; use original docx as reference to keep styles
        out_path = td / "out.docx"
        pypandoc.convert_text(
            md,
            to="docx",
            format="markdown+tex_math_dollars",
            outputfile=str(out_path),
            extra_args=["--reference-doc", str(in_path)],
        )
        return out_path.read_bytes()


# 8) UI: sidebar
# ============================================================
ocr_timeout_default, translate_timeout_default = get_timeout_defaults()

def _init_state():
    defaults = {
        "ocr_model_id": get_default_model(),
        "translate_model_id": os.environ.get("ARK_TRANSLATE_MODEL") or DEFAULT_TRANSLATE_MODEL,
        "mml2omml_xsl": os.environ.get(DEFAULT_MML2OMML_XSL_ENV, ""),
        "max_side": 2200,
        "tile_h": 1600,
        "overlap": 160,
        "jpeg_q": 85,
        "out_tokens": 4096,
        "ocr_timeout_s": int(ocr_timeout_default),
        "translate_timeout_s": int(translate_timeout_default),
        "preset_mode": "Balanced",
        "glossary_allow_substring": False,
        "smart_glossary_top_k": 20,
        "smart_glossary_preview_chars": 9000,
        "smart_glossary_candidates": [],
        # MERGE NOTE: keep all 3 keys below (text/editor/display) when resolving branch conflicts.
        "formula_latex_text": "",
        "formula_latex_editor": "",
        "formula_display_mode": "行间（$$...$$）",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

_init_state()

# ✅ 关键修复：在任何 widget 创建之前应用 pending 推荐参数
if "__pending_rec__" in st.session_state:
    rec = st.session_state.pop("__pending_rec__")
    for k in ["max_side", "tile_h", "overlap", "jpeg_q", "out_tokens"]:
        if k in rec:
            st.session_state[k] = int(rec[k])
    st.session_state["__pending_rec_msg__"] = f"已回填推荐参数：{rec}"

with st.sidebar:
    st.subheader("Ark 配置")
    st.text_input("OCR model（默认 EP 已填）", key="ocr_model_id")
    st.text_input("Translate model（翻译专用 EP）", key="translate_model_id")
    st.text_input("MML2OMML.XSL 路径（可选：仅用于把 $...$ 转 Word 公式）", key="mml2omml_xsl")
    st.caption(f"Base URL: {get_ark_base_url()}")

    st.divider()

    with st.expander("术语表/一致性 & 译后质量检查（自动重试）", expanded=False):
        st.checkbox("启用译后质量检查（自动重试）", key="qa_enable", value=True)
        st.slider("自动重试次数（仅在检测到问题时）", min_value=0, max_value=5, value=2, step=1, key="qa_retries")
        st.text_area(
            "术语表（每行：源语 => 译语 / 源语\t译语 / 源语->译语）",
            key="glossary_text",
            height=160,
            help="用于保证术语一致性：翻译时会把源术语替换为占位符，译后再回填为指定译语。",
        )
        st.checkbox(
            "术语允许子串匹配（默认关闭，建议仅在明确需要时开启）",
            key="glossary_allow_substring",
            value=False,
            help="关闭时：英文/数字术语按词边界匹配，符号术语按前后非字母数字边界匹配，中文走最长优先匹配。",
        )
        st.caption("提示：默认关闭可避免 cat→catastrophe、AI→PAIN 这类污染替换。")
    st.subheader("OCR 参数与自动推荐")

    st.selectbox("预设", ["Fast", "Balanced", "Accurate"], key="preset_mode",
                 help="Fast：更快更省；Accurate：更清晰更稳但更慢；Balanced：折中。")

    st.slider("Max side（最长边像素）", 800, 3200, key="max_side", step=100,
              help="大：更清晰更准但更慢更贵；小：更快更省但小字/公式易错。")
    st.slider("Tile height（切片高度）", 800, 2600, key="tile_h", step=100,
              help="大：切片少更快但易截断；小：切片多更稳但更慢更贵。")
    st.slider("Overlap（切片重叠）", 0, 400, key="overlap", step=10,
              help="大：边界漏字更少但可能重复；小：更快但边界更易漏。")
    st.slider("JPEG quality（压缩质量）", 50, 95, key="jpeg_q", step=1,
              help="高：细节更清晰更准但更慢；低：更快但公式/小字更易糊。")
    st.slider("OCR max tokens（输出上限）", 1024, 8192, key="out_tokens", step=256,
              help="输出太短会截断漏字。截断优先：减 tile_h；其次：增 tokens。")

    st.number_input("OCR 请求超时（秒）", min_value=30, max_value=600, key="ocr_timeout_s", step=10,
                    help="大图/网络慢可调到 180-300。")
    st.number_input("翻译请求超时（秒）", min_value=30, max_value=900, key="translate_timeout_s", step=10,
                    help="10页以上建议 180-300；更稳可更高。")

    with st.expander("参数速查（给不懂的人）", expanded=False):
        st.markdown(
            """
- **输出经常断在半页**：先把 **Tile height** 调小（1600→1200），再把 **OCR max tokens** 调大（4096→6144）。
- **速度太慢/切片太多**：把 **Tile height** 调大（1600→2000），再把 **Max side** 略降（2200→1800）。
- **小字/公式错多**：把 **Max side** 调大（2200→2800+），或 **JPEG** 85→90/95。
- **翻译大文档更稳**：把 **max_batch_chars** 调小（12000→8000/6000），同时把超时调高（180→300）。
"""
        )


# ============================================================
# 9) Tabs
# ============================================================
# MERGE NOTE: keep Tab⑤ label and corresponding with tabs[4] block together during conflict resolution.
tabs = st.tabs([
    "① PDF/图片 OCR → 导出",
    "② 公式 OCR → LaTeX（沉浸编辑）",
    "③ Word(.docx) → 保排版翻译/就地替换公式（best-effort）",
    "④ Word(.docx) → LaTeX/Markdown 直接导出（推荐）",
    "⑤ 标题格式化",
])

# ---------------------------
# Tab 1: OCR (PDF + images)
# ---------------------------
with tabs[0]:
    st.subheader("PDF/图片 OCR（带自动参数推荐 + 进度条）")
    st.write("建议：先上传 1 页代表性页面 → 点“自动推荐参数” → 再批量处理。")

    if "__pending_rec_msg__" in st.session_state:
        st.success(st.session_state.pop("__pending_rec_msg__"))

    files = st.file_uploader(
        "上传 PDF 或图片（可多选，PDF 将按页渲染）",
        type=["pdf", "png", "jpg", "jpeg", "webp"],
        accept_multiple_files=True
    )

    col_pdf = st.columns([1, 1, 2])
    with col_pdf[0]:
        pdf_dpi = st.selectbox("PDF 渲染 DPI", [200, 300, 400], index=1)
    with col_pdf[1]:
        pdf_max_pages = st.number_input("PDF 最多页数（0=不限制）", min_value=0, max_value=500, value=0, step=1)
        pdf_page_range = st.text_input("PDF 页范围（可选）", value="", help="例如 1-3,5,8-10；留空=按‘最多页数’或全文")
    with col_pdf[2]:
        if not HAVE_PYMUPDF:
            st.warning("未检测到 pymupdf，PDF 上传不可用。请 pip install pymupdf")

    # 先把所有输入转换成 images list
    images: List[Image.Image] = []
    page_meta: List[str] = []  # for page labeling

    if files:
        for f in files:
            name = getattr(f, "name", "upload")
            ext = (name.split(".")[-1] or "").lower()
            if ext == "pdf":
                if not HAVE_PYMUPDF:
                    st.error("你上传了 PDF，但环境缺少 pymupdf。请安装后重试：pip install pymupdf")
                    st.stop()
                pdf_bytes = read_uploaded_bytes(f)
                try:
                    imgs = pdf_bytes_to_images(
                        pdf_bytes,
                        dpi=int(pdf_dpi),
                        max_pages=None if int(pdf_max_pages) == 0 else int(pdf_max_pages),
                        page_indices=parse_page_range(pdf_page_range),
                    )
                except Exception as e:
                    st.error(f"PDF 渲染失败：{e}")
                    st.stop()
                for i, im in enumerate(imgs, start=1):
                    images.append(im)
                    page_meta.append(f"{name} - p{i}")
            else:
                try:
                    im = Image.open(f)
                    images.append(im)
                    page_meta.append(name)
                except Exception as e:
                    st.warning(f"无法读取图片 {name}：{e}")

    cols = st.columns([1, 1, 2])
    with cols[0]:
        auto_btn = st.button("🪄 自动推荐参数（基于第1页）", disabled=not images)
    with cols[1]:
        preset_btn = st.button("🎛️ 应用预设（Fast/Balanced/Accurate）", disabled=not images)
    with cols[2]:
        st.caption("自动推荐会根据图片高度与文字/边缘密度估计调整 max_side/tile/overlap/tokens。")

    if images and auto_btn:
        rec = recommend_ocr_params(images[0], mode=st.session_state["preset_mode"])
        st.session_state["__pending_rec__"] = rec
        st.rerun()

    if images and preset_btn:
        rec = recommend_ocr_params(images[0], mode=st.session_state["preset_mode"])
        st.session_state["__pending_rec__"] = rec
        st.rerun()

    join_lines = st.checkbox("可选：合并断行（适合 PDF 每行强制换行）", value=False)

    def merge_hard_wraps(md: str) -> str:
        parts = md.split("\n\n")
        out = []
        for p in parts:
            lines = [x.strip() for x in p.splitlines()]
            if any(l.startswith(("-", "*", "|", "#")) for l in lines):
                out.append("\n".join(p.splitlines()))
            else:
                out.append(" ".join([l for l in lines if l != ""]).strip())
        return "\n\n".join(out).strip()

    if st.button("开始 OCR 并导出", type="primary", disabled=not images):
        if not st.session_state["ocr_model_id"].strip():
            st.error("请先填写 model（ep-xxxx 或模型ID）。")
            st.stop()

        client = get_ark_client(default_timeout_s=int(st.session_state["ocr_timeout_s"]))

        pages = []
        total_pages = len(images)
        page_bar = st.progress(0, text="准备 OCR…")

        for pi, img in enumerate(images, start=1):
            tile_bar = st.progress(0, text=f"OCR 第 {pi}/{total_pages} 页：准备切片…")

            def _tile_cb(cur: int, total: int):
                tile_bar.progress(int(cur / total * 100), text=f"OCR 第 {pi}/{total_pages} 页：切片 {cur}/{total}")

            md = ocr_image_to_markdown(
                client=client,
                model=st.session_state["ocr_model_id"].strip(),
                img=img,
                max_side=int(st.session_state["max_side"]),
                tile_h=int(st.session_state["tile_h"]),
                overlap=int(st.session_state["overlap"]),
                jpeg_q=int(st.session_state["jpeg_q"]),
                out_tokens=int(st.session_state["out_tokens"]),
                timeout_s=int(st.session_state["ocr_timeout_s"]),
                progress_cb=_tile_cb,
            )

            if join_lines:
                md = merge_hard_wraps(md)

            label = page_meta[pi - 1] if pi - 1 < len(page_meta) else f"Page {pi}"
            pages.append(f"## 第 {pi} 页（{label}）\n\n{md}")
            page_bar.progress(int(pi / total_pages * 100), text=f"已完成 {pi}/{total_pages} 页")
            tile_bar.empty()

        merged_md = normalize_md("\n\n---\n\n".join(pages))

        st.success("OCR 完成")
        st.markdown("### 预览（Markdown）")
        st.code(merged_md, language="markdown")

        v1_docx = pandoc_md_to_docx(merged_md)
        v2_md_bytes = merged_md.encode("utf-8")

        v3_md = md_to_latex_code_style(merged_md)
        v3_docx = pandoc_md_to_docx(v3_md)

        st.download_button("下载 V1：Rendered.docx（pandoc渲染公式）", data=v1_docx, file_name="OCR_Rendered.docx")
        st.download_button("下载 V2：Result.md（原始 Markdown）", data=v2_md_bytes, file_name="OCR_Result.md")
        st.download_button("下载 V3：LaTeX_equation_code.docx（公式为 LaTeX 代码块）", data=v3_docx,
                           file_name="OCR_LaTeX_equation_code.docx")
        st.download_button("下载 V3：LaTeX_equation_code.md", data=v3_md.encode("utf-8"),
                           file_name="OCR_LaTeX_equation_code.md")


# ---------------------------
# Tab 2: Formula OCR -> LaTeX immersive editor
# ---------------------------
with tabs[1]:
    st.subheader("公式 OCR → LaTeX（可编辑 + 实时预览）")
    st.caption("上传公式截图，自动转 LaTeX；可手动编辑并实时预览，支持复制。")

    formula_file = st.file_uploader("上传公式图片", type=["png", "jpg", "jpeg", "webp"], key="formula_ocr_img")
    cfm1, cfm2, cfm3 = st.columns([1, 1, 2])
    with cfm1:
        formula_out_tokens = st.slider("公式 OCR max tokens", 256, 4096, 1024, 128, key="formula_ocr_tokens")
    with cfm2:
        st.selectbox("预览模式", ["行间（$$...$$）", "行内（$...$）"], key="formula_display_mode")
    with cfm3:
        run_formula_ocr = st.button("识别公式并填入编辑器", type="primary", disabled=not formula_file, key="btn_formula_ocr")

    if run_formula_ocr and formula_file:
        try:
            client = get_ark_client(default_timeout_s=int(st.session_state["ocr_timeout_s"]))
            img = Image.open(io.BytesIO(read_uploaded_bytes(formula_file))).convert("RGB")
            data_url = _jpeg_bytes_to_data_url(pil_to_jpeg_bytes(img, quality=int(st.session_state.get("jpeg_q", 90))))
            messages = [{
                "role": "user",
                "content": [
                    {"type": "text", "text": FORMULA_OCR_TO_LATEX_PROMPT_ZH},
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }]
            res = safe_chat_completions(
                client=client,
                model=st.session_state.get("ocr_model_id", "").strip() or get_default_model(),
                messages=messages,
                temperature=0.0,
                max_tokens=int(formula_out_tokens),
                timeout_s=int(st.session_state["ocr_timeout_s"]),
                retries=6,
            )
            if res.error_message:
                st.error(f"公式 OCR 失败：{res.error_message}")
            else:
                _ocr_latex = (res.text or "").strip()
                st.session_state["formula_latex_text"] = _ocr_latex
                st.session_state["formula_latex_editor"] = _ocr_latex
                st.session_state["formula_latex_text"] = (res.text or "").strip()
                st.success("公式识别完成，可在下方继续编辑。")
        except Exception as e:
            st.error(f"公式 OCR 调用失败：{type(e).__name__}: {e}")
    # 关键修复：公式编辑器只创建一次，key 全局唯一。
    # 识别结果写入 st.session_state["formula_latex_editor"]，并作为 text_area 的 value。
    latex_text = st.text_area(
        "LaTeX 编辑器",
        value=st.session_state.get("formula_latex_editor", st.session_state.get("formula_latex_text", "")),
        height=220,
        key="formula_latex_editor",
    )
    st.session_state["formula_latex_text"] = latex_text
    st.session_state["formula_latex_editor"] = latex_text



    if latex_text.strip():
        preview_expr = latex_text.strip()
        st.markdown("**实时预览**")
        try:
            if st.session_state.get("formula_display_mode") == "行内（$...$）":
                st.markdown(f"预览：${preview_expr}$")
            else:
                st.latex(preview_expr)
        except Exception as e:
            st.warning(f"预览渲染失败：{type(e).__name__}: {e}")

    wrapped = f"${latex_text.strip()}$" if st.session_state.get("formula_display_mode") == "行内（$...$）" else f"$$\n{latex_text.strip()}\n$$"
    st.code(wrapped if latex_text.strip() else "", language="latex")
    st.download_button(
        "下载 LaTeX 文本（.txt）",
        data=(wrapped if latex_text.strip() else "").encode("utf-8"),
        file_name="formula_latex.txt",
        mime="text/plain",
        key="dl_formula_latex_txt",
    )


# ---------------------------
# Tab 3: DOCX translate + best-effort equation replace
# ---------------------------
with tabs[2]:
    st.subheader("Word(.docx) → 保排版翻译 / 公式就地替换（best-effort）")
    st.warning(
        "说明：Word 可编辑公式（OMML）要“可靠”转 LaTeX，推荐使用 Tab③ 的 Pandoc 直接导出。\n"
        "本 Tab 的“就地替换公式”为 best-effort，可能出现错位/不全。"
    )

    docx_file = st.file_uploader("上传 Word 文档（.docx）", type=["docx"], key="docx_inplace")

    colA, colB, colC = st.columns([1, 1, 1])
    with colA:
        do_equation_replace = st.checkbox("把 Word 原生公式（OMML）替换为 LaTeX 代码（best-effort）", value=False)
        do_latex_to_omml = st.checkbox("检测 $...$ / $$...$$ 并转换为 Word 公式（OMML）", value=False)
        do_translate = st.checkbox("翻译文档（尽量保持原排版）", value=False)

    with colB:
        src_lang = st.selectbox("源语言", ["Auto", "Chinese", "English", "Japanese", "Korean", "Spanish"], index=0)
        dst_lang = st.selectbox("目标语言", ["Chinese", "English", "Japanese", "Korean", "Spanish"], index=1)

    with colC:
        auto_batch = st.checkbox("自动分批（推荐）", value=True)
        target_batches = st.number_input("目标批次数（自动分批用）", min_value=1, max_value=50, value=8, step=1)

    if not auto_batch:
        max_batch_chars = st.slider("翻译分批大小（max_batch_chars）", 4000, 20000, 12000, 500)
    else:
        max_batch_chars = 12000  # will be computed after parsing doc

    show_translate_diagnostics = st.checkbox("显示翻译抓取诊断（建议打开排错）", value=True)

    # ---------------------------
    # 翻译增强：Pandoc 中转（OCR $$ 风格）→ 可选导出 Markdown / Word(OMML) / LaTeX
    # ---------------------------
    st.divider()
    st.subheader("翻译增强：Pandoc 中转（OCR $$ 风格）→ 可选导出 Markdown / Word(OMML) / LaTeX")
    st.markdown(
        '<div class="hint-card">'
        "这个模式参考 Tab① 的 OCR 输出风格：先用 Pandoc 把 Word 内容抽成 Markdown，并把公式统一为 $...$ / $$...$$；"
        "再对文本做翻译（公式占位符严格保留），最后可直接导出："
        "<b>译文 Markdown（OCR 风格）</b>、<b>译文 Word（Pandoc 写回，公式为可编辑 OMML）</b>、<b>译文 LaTeX</b>。"
        "<br/>适合：你更看重“翻译覆盖率/稳定性”，而不是 100% 保留 Word 原排版。"
        "</div>",
        unsafe_allow_html=True,
    )

    enable_pandoc_translate = st.checkbox("启用 Pandoc 中转翻译（更稳，推荐）", value=True, key="pandoc_translate_enable")
    max_batch_chars2 = st.number_input("每批最大字符数（自动分段）", min_value=2000, max_value=20000, value=8000, step=500, key="pandoc_translate_max_chars")

    with st.expander("智能术语提取（Smart Glossary Extraction）", expanded=False):
        st.caption("翻译前可先预扫描文档前部内容，自动提取 top-K 术语候选，勾选后写入术语表。")
        csg1, csg2 = st.columns(2)
        with csg1:
            st.number_input("预扫描字符数", min_value=2000, max_value=30000, step=500, key="smart_glossary_preview_chars")
        with csg2:
            st.number_input("候选数量 Top-K", min_value=5, max_value=100, step=1, key="smart_glossary_top_k")

        if st.button("预扫描并提取术语候选", disabled=not docx_file, key="btn_smart_glossary_scan"):
            try:
                doc_preview_bytes = read_uploaded_bytes(docx_file)
                preview_text = mp.extract_docx_preview_text_for_glossary(
                    doc_preview_bytes,
                    max_chars=int(st.session_state.get("smart_glossary_preview_chars", 9000)),
                )
                client = get_ark_client(default_timeout_s=int(st.session_state["translate_timeout_s"]))
                candidates = smart_extract_glossary_candidates(
                    client=client,
                    model=st.session_state.get("translate_model_id", "").strip() or DEFAULT_TRANSLATE_MODEL,
                    preview_text=preview_text,
                    src_lang=src_lang,
                    dst_lang=dst_lang,
                    top_k=int(st.session_state.get("smart_glossary_top_k", 20)),
                    timeout_s=int(st.session_state["translate_timeout_s"]),
                )
                st.session_state["smart_glossary_candidates"] = candidates
                st.success(f"已生成术语候选：{len(candidates)} 条")
            except Exception as e:
                st.error(f"术语提取失败：{type(e).__name__}: {e}")

        candidates = st.session_state.get("smart_glossary_candidates", []) or []
        if candidates:
            st.markdown("**候选（可勾选后写入术语表）**")
            for idx, it in enumerate(candidates):
                src_v = it.get("src", "")
                sug_v = it.get("suggested", "")
                type_v = it.get("type", "term")
                freq_v = it.get("freq_estimate", 1)
                note_v = it.get("note", "")
                cols = st.columns([1, 3, 3, 1, 3])
                checked = cols[0].checkbox("选", key=f"sg_pick_{idx}", value=True)
                src_edit = cols[1].text_input("src", value=src_v, key=f"sg_src_{idx}")
                sug_edit = cols[2].text_input("suggested", value=sug_v, key=f"sg_sug_{idx}")
                cols[3].write(f"{type_v}/{freq_v}")
                cols[4].caption(note_v or "-")
                it["_picked"] = bool(checked)
                it["src"] = src_edit.strip()
                it["suggested"] = sug_edit.strip()

            if st.button("将已勾选候选写入术语表", key="btn_apply_smart_glossary"):
                picked_lines: List[str] = []
                for it in candidates:
                    if it.get("_picked") and it.get("src") and it.get("suggested"):
                        picked_lines.append(f"{it['src']} => {it['suggested']}")
                merged = (st.session_state.get("glossary_text", "") or "").strip()
                addon = "\n".join(picked_lines).strip()
                st.session_state["glossary_text"] = (merged + "\n" + addon).strip() if merged and addon else (merged or addon)
                st.success(f"已写入术语表：{len(picked_lines)} 条。提示：英文术语建议同时补充大小写/复数变体。")

    if st.button("开始翻译（Pandoc 中转）", type="primary", disabled=(not docx_file) or (not enable_pandoc_translate), key="btn_pandoc_translate"):
        try:
            doc_bytes2 = read_uploaded_bytes(docx_file)
            sha2 = mp.sha256_bytes(doc_bytes2)
            with st.spinner("Pandoc 中转翻译中（会自动缓存，重复操作不重复付费）..."):
                md_tr, out_docx_tr, latex_tr, stats = _cached_pandoc_translate_ocr_route(
                    sha2,
                    doc_bytes2,
                    src_lang=src_lang,
                    dst_lang=dst_lang,
                    timeout_s=int(st.session_state["translate_timeout_s"]),
                    max_batch_chars=int(max_batch_chars2),
                    glossary_text=st.session_state.get("glossary_text", ""),
                    glossary_allow_substring=bool(st.session_state.get("glossary_allow_substring", False)),
                )

            st.success("Pandoc 翻译完成")
            st.session_state["tab3_seed_markdown"] = md_tr  # 传给 Tab③ 进一步导出

            c1, c2, c3 = st.columns(3)
            with c1:
                st.download_button("下载：Translated_OCRStyle.md", data=md_tr.encode("utf-8"), file_name="Translated_OCRStyle.md", mime="text/markdown", key="dl_tr_md")
            with c2:
                st.download_button("下载：Translated_OMML.docx", data=out_docx_tr, file_name="Translated_OMML.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_tr_docx")
            with c3:
                st.download_button("下载：Translated.tex", data=latex_tr.encode("utf-8"), file_name="Translated.tex", mime="text/x-tex", key="dl_tr_tex")

            with st.expander("翻译统计 / 排错", expanded=False):
                st.json(stats, expanded=True)

        except Exception as e:
            st.error(f"Pandoc 翻译失败：{type(e).__name__}: {e}")



    if st.button("开始处理并导出（new.docx）", type="primary", disabled=not docx_file):
        if not st.session_state["ocr_model_id"].strip():
            st.error("请先在侧边栏配置 OCR model（用于 OCR/PDF，不影响翻译）。")
            st.stop()

        # 翻译不依赖侧边栏输入的 translate_model_id（会强制使用 DEFAULT_TRANSLATE_MODEL），
        # 但这里仍保留输入框，便于你看到当前配置。
        client = get_ark_client(default_timeout_s=int(st.session_state["translate_timeout_s"]))

        doc_bytes = read_uploaded_bytes(docx_file)
        doc = Document(io.BytesIO(doc_bytes))

        # 0) 将文本中的 LaTeX 公式 ($...$ / $$...$$ / \(\) / \[\]) 转成 Word 原生公式（可选）
        if do_latex_to_omml:
            if not HAVE_LATEX_OMML:
                st.error("该功能需要额外依赖：lxml + latex2mathml（见 requirements 更新）。")
                st.stop()

            xsl_cfg = (st.session_state.get("mml2omml_xsl", "") or "").strip()
            xsl_env = (os.environ.get(DEFAULT_MML2OMML_XSL_ENV, "") or "").strip()
            xsl_path = xsl_cfg or xsl_env

            if not xsl_path:
                st.warning("未提供 MML2OMML.XSL 路径：已跳过 LaTeX→OMML（不影响翻译/公式导出）。")
            else:
                try:
                    converted = convert_inline_latex_to_omml_in_doc(
                        doc,
                        xsl_path=xsl_path,
                        log=lambda s: st.info(s),
                    )
                    st.success(f"LaTeX→OMML 转换完成：{converted} 处")
                except Exception as e:
                    st.warning(f"LaTeX→OMML 转换失败（已跳过，不影响后续）：{type(e).__name__}: {e}")

        # 1) 把 Word 原生公式（OMML）替换为 Pandoc 可渲染的 LaTeX（$...$ / $$...$$）
        if do_equation_replace:
            with st.spinner("提取公式序列（pandoc）..."):
                math_seq = extract_math_from_docx_with_pandoc(doc_bytes)
            with st.spinner("替换 Word 原生公式为 LaTeX 代码（best-effort）..."):
                replaced_count = replace_omml_with_latex_code(doc, math_seq, use_equation_env=False)
            st.info(f"已替换公式数量（best-effort）：{replaced_count}")

        # 2) 翻译（保持原排版：仅替换 runs 的文本；图片/表格结构不动）
        if do_translate:
            batch_bar = st.progress(0, text="准备翻译…")
            diag_box = st.empty()
            attempt_box = st.empty()

            def _batch_cb(cur: int, total: int):
                batch_bar.progress(int(cur / total * 100), text=f"翻译批次 {cur}/{total}")

            diag_payload_holder: Dict[str, Any] = {}

            def _diag_cb(payload: Dict[str, Any]):
                diag_payload_holder.update(payload)

            # 预扫描一次 items 用于“自动分批”估算
            all_ps = iter_all_paragraphs_extended(doc)
            items_preview: List[dict] = []
            pid = 0
            for p in all_ps:
                runs = [r for r in p.runs if r.text is not None and r.text != ""]
                if not runs:
                    continue
                segs = []
                for r in runs:
                    protected, _ = protect_math(r.text)
                    segs.append(protected)
                if not "".join(segs).strip():
                    continue
                pid += 1
                items_preview.append({"id": f"p{pid}", "segments": segs})

            if auto_batch:
                max_batch_chars_local = estimate_auto_batch_chars(items_preview, target_batches=int(target_batches))
                st.info(f"自动分批：items={len(items_preview)}，估算 max_batch_chars={max_batch_chars_local}")
            else:
                max_batch_chars_local = int(max_batch_chars)

            if show_translate_diagnostics:
                diag_box.info(
                    "诊断将在开始后显示：扫描段落数 / 可翻译 items 数 / runs 数 / 字符数。\n"
                    "若 items=0，多半是内容在文本框/形状里（python-docx 读不到）。"
                )

            with st.spinner("翻译中（分批提交，保持图片/公式位置不动）..."):
                translate_docx_in_place(
                    doc=doc,
                    client=client,
                    model=st.session_state.get("translate_model_id", "").strip(),
                    src_lang=src_lang,
                    dst_lang=dst_lang,
                    max_batch_chars=int(max_batch_chars_local),
                    timeout_s=int(st.session_state["translate_timeout_s"]),
                    progress_cb=_batch_cb,
                    diagnostic_cb=_diag_cb,
                    attempt_msg_cb=(lambda msg: attempt_box.info(msg)),
                )

            if show_translate_diagnostics and diag_payload_holder:
                diag_box.success(
                    f"翻译抓取诊断：\n"
                    f"- paragraphs_scanned = {diag_payload_holder.get('paragraphs_scanned')}\n"
                    f"- items_built       = {diag_payload_holder.get('items_built')}\n"
                    f"- runs_counted      = {diag_payload_holder.get('runs_counted')}\n"
                    f"- visible_chars     = {diag_payload_holder.get('visible_chars_counted')}\n"
                )

        out_docx_bytes = doc_to_bytes(doc)
        st.success("处理完成")
        st.download_button("下载 new.docx", data=out_docx_bytes, file_name="new.docx")



# ---------------------------
# Tab 3: DOCX -> LaTeX/Markdown export (recommended)
# ---------------------------
with tabs[3]:
    st.markdown("---")
    st.subheader("Word(.docx) → LaTeX / Markdown 直接导出（推荐：可编辑公式最稳）")
    st.write("这个模式不追求保留 Word 排版，而追求“学术 LaTeX 输出正确性”，尤其适合大量可编辑公式。")

    # 你可以在 Tab②（Pandoc 中转翻译）里把译文 Markdown 送到这里进一步导出
    seed_md = st.session_state.get("tab3_seed_markdown", "")
    if seed_md:
        with st.expander("Tab② 传入的译文 Markdown（OCR $$ 风格）→ 直接导出", expanded=False):
            md_in = st.text_area("Markdown 输入", value=seed_md, height=220, key="tab3_seed_md_area")
            colx1, colx2 = st.columns(2)
            with colx1:
                if st.button("生成 Word（Pandoc 写回，公式 OMML）", key="tab3_seed_to_docx"):
                    try:
                        out_docx = mp.pandoc_markdown_to_docx(md_in)
                        st.download_button(
                            "下载：SeedTranslated_OMML.docx",
                            data=out_docx,
                            file_name="SeedTranslated_OMML.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    except Exception as e:
                        st.error(f"生成 DOCX 失败：{type(e).__name__}: {e}")
            with colx2:
                if st.button("生成 LaTeX（.tex）", key="tab3_seed_to_tex"):
                    try:
                        tex = mp.pandoc_markdown_to_latex(md_in)
                        st.download_button("下载：SeedTranslated.tex", data=tex.encode("utf-8"), file_name="SeedTranslated.tex", mime="text/x-tex")
                    except Exception as e:
                        st.error(f"生成 LaTeX 失败：{type(e).__name__}: {e}")



    docx_file2 = st.file_uploader("上传 Word 文档（.docx）", type=["docx"], key="docx_export")

    out_format = st.selectbox("导出格式", ["latex (.tex)", "markdown (.md)"], index=0)
    wrap_none = st.checkbox("wrap=none（不自动换行）", value=True)

    # ✅ 新增：把 pandoc 默认的 \[...\]/\(...\) 统一成更“论文友好”的格式
    st.markdown("**公式输出格式（可选）**")
    col_m1, col_m2 = st.columns([1, 1])
    with col_m1:
        display_style = st.selectbox(
            "行间公式（display math）",
            ["equation（带编号）", "equation*（不编号）", "bracket（\\[...\\]）", "dollars（$$...$$）"],
            index=0,
            help="pandoc 默认常用 \\[...\\]。论文通常更喜欢 equation/equation*，便于统一编号与引用。",
        )
    with col_m2:
        inline_style = st.selectbox(
            "行内公式（inline math）",
            ["paren（\\(...\\)）", "dollars（$...$）"],
            index=0,
            help="pandoc 默认常用 \\(...\\)。如果你后续要在 Markdown 里编辑，可能更喜欢 $...$。",
        )

    # 映射到 normalize_pandoc_math 的参数值
    display_style_key = display_style.split("（", 1)[0].strip()
    inline_style_key = inline_style.split("（", 1)[0].strip()

    if st.button("导出", type="primary", disabled=not docx_file2):
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            in_path = td / "in.docx"
            in_path.write_bytes(docx_file2.read())

            extra_args = []
            if wrap_none:
                extra_args += ["--wrap=none"]

            if out_format.startswith("latex"):
                # docx -> latex
                out_text = pypandoc.convert_file(
                    str(in_path),
                    to="latex",
                    format="docx",
                    extra_args=extra_args,
                )
                # 统一公式风格（\[\] / $$ $$ / \(\) / $ $）
                out_text = normalize_pandoc_math(
                    out_text,
                    display_style=display_style_key,
                    inline_style=("dollars" if inline_style_key.startswith("dollars") else "paren"),
                )

                st.code(out_text[:4000] + ("\n...\n" if len(out_text) > 4000 else ""), language="latex")
                st.download_button("下载 .tex", data=out_text.encode("utf-8"), file_name="export.tex")
            else:
                # docx -> markdown
                # 关键：显式打开 tex_math_dollars 扩展，保证 $...$ / $$...$$ 被当作数学处理
                # Pandoc 支持通过 format name +EXTENSION 来启用扩展（见官方手册）。
                out_text = pypandoc.convert_file(
                    str(in_path),
                    to="markdown+tex_math_dollars",
                    format="docx",
                    extra_args=extra_args,
                )

                # Markdown 输出也可按同一套规则统一（通常建议保留 dollars 方便后续编辑）
                out_text = normalize_pandoc_math(
                    out_text,
                    display_style=display_style_key,
                    inline_style=("dollars" if inline_style_key.startswith("dollars") else "paren"),
                )

                st.code(out_text[:4000] + ("\n...\n" if len(out_text) > 4000 else ""), language="markdown")
                st.download_button("下载 .md", data=out_text.encode("utf-8"), file_name="export.md")



    st.markdown("---")
    st.subheader("把文档中的 ... / ... 文本公式转为可编辑 Word 公式（OMML）")
    st.caption("适用场景：你在 Word 里手打了 $...$（或 $$...$$）作为占位公式，想一键转成可编辑的原生公式对象（OMML）。实现方式：docx → pandoc markdown → docx（用原文档作为 reference-doc 尽量保留样式）。")
    docx_file3 = st.file_uploader("上传 Word 文档（.docx）", type=["docx"], key="docx_omml_roundtrip")
    if st.button("生成：EditableEquations.docx", key="btn_omml_roundtrip", disabled=not docx_file3):
        try:
            # 统一 getvalue() 读取，避免 UploadedFile.read() 二次读取为空
            in_bytes = read_uploaded_bytes(docx_file3)
            out_docx_bytes = docx_roundtrip_make_equations_editable(in_bytes)
            st.download_button(
                "下载 EditableEquations.docx",
                data=out_docx_bytes,
                file_name="EditableEquations.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_omml_roundtrip",
            )
            st.success("已生成：EditableEquations.docx（公式已转为可编辑 OMML）")
        except Exception as e:
            st.error(f"生成失败：{type(e).__name__}: {e}")


    st.markdown("---")
    st.subheader("AI 增强：DOCX 文本 → 类似图片 OCR 的公式风格 → 再转 OMML / Pandoc")
    st.caption(
        "思路：先用 Pandoc 把 docx 里的公式变成 $...$ / $$...$$ 文本，再交给模型做“只改公式不改正文”的校正，"
        "最后再用 Pandoc 写回 Word（OMML）。这能显著减少少量顽固公式的格式异常。"
    )

    enable_ai_fix = st.checkbox("开启 AI 公式校正（更慢但更准）", value=False, key="ai_fix_enable")

    col_ai1, col_ai2 = st.columns([2, 1])
    with col_ai1:
        ai_model = st.text_input(
            "AI 校正使用的 model（默认用 OCR model）",
            value=st.session_state.get("ocr_model_id", "").strip(),
            key="ai_fix_model",
            disabled=not enable_ai_fix,
        )
        max_batch_chars = st.number_input(
            "AI 每次输入最大字符数（自动分段）",
            min_value=2000,
            max_value=30000,
            value=8000,
            step=500,
            disabled=not enable_ai_fix,
        )
    with col_ai2:
        ai_out_tokens = st.number_input(
            "AI 输出 tokens 上限",
            min_value=512,
            max_value=8192,
            value=int(st.session_state.get("out_tokens", 4096)),
            step=256,
            disabled=not enable_ai_fix,
        )
        ai_timeout_s = st.number_input(
            "AI 单次请求超时（秒）",
            min_value=30,
            max_value=600,
            value=int(st.session_state.get("ocr_timeout_s", 120)),
            step=10,
            disabled=not enable_ai_fix,
        )

    docx_file_ai = st.file_uploader("上传 Word 文档（.docx）", type=["docx"], key="docx_ai_roundtrip")

    if st.button("生成：AI_EditableEquations.docx", key="btn_ai_roundtrip", disabled=not (enable_ai_fix and docx_file_ai)):
        try:
            if not ai_model.strip():
                st.error("AI model 为空：请在侧边栏填写 OCR model，或在此处手动填写。")
                st.stop()

            in_bytes = read_uploaded_bytes(docx_file_ai)
            if not in_bytes:
                st.error("上传文件为空（可能重复点击导致 read() 被消费）。请重新上传或刷新页面。")
                st.stop()

            sha = mp.sha256_bytes(in_bytes)

            with st.spinner("AI 校正 + Pandoc 回写中…（首次可能较慢，后续同文件会走缓存）"):
                out_docx_bytes, md_used = _cached_ai_roundtrip(
                    sha,
                    in_bytes,
                    model=ai_model.strip(),
                    max_batch_chars=int(max_batch_chars),
                    timeout_s=int(ai_timeout_s),
                    out_tokens=int(ai_out_tokens),
                    base_url=str(get_ark_base_url()),
                )

            st.download_button(
                "下载 AI_EditableEquations.docx",
                data=out_docx_bytes,
                file_name="AI_EditableEquations.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_ai_roundtrip",
            )
            with st.expander("查看 AI 校正后用于 Pandoc 的 Markdown（调试用）", expanded=False):
                st.code((md_used or "")[:12000] + ("\n...\n" if (md_used and len(md_used) > 12000) else ""), language="markdown")

            st.success("已生成：AI_EditableEquations.docx")

        except Exception as e:
            st.error(f"生成失败：{type(e).__name__}: {e}")



# ============================================================
# 9) Cache-enabled wrappers (keep old call sites unchanged)
# ============================================================
def docx_to_pandoc_markdown_for_math(docx_bytes: bytes) -> str:
    sha = mp.sha256_bytes(docx_bytes)
    return _cached_docx_to_md_for_math(sha, docx_bytes)



# ---------------------------
# Tab 5: Title formatting (python-docx only)
# ---------------------------
if len(tabs) > 4:
    with tabs[4]:
        st.markdown("---")
        st.subheader("标题格式化（纯程序逻辑，无 AI）")
        st.caption("上传 .docx 后，按段落样式识别 Heading 1/2/3，并按你配置的字体名/字号/粗体批量写回导出。")

        st.markdown(
            "**说明（识别规则）**：仅根据段落样式名判断标题级别（如 `Heading 1/2/3`、`标题 1/2/3`），其余视为普通段落。"
        )

        docx_title_file = st.file_uploader("上传 Word 文档（.docx）", type=["docx"], key="docx_title_format")

        col_t1, col_t2, col_t3 = st.columns(3)
        with col_t1:
            st.markdown("**一级标题样式**")
            h1_font = st.text_input("字体名（H1）", value="SimHei", key="h1_font")
            h1_size = st.number_input("字号 pt（H1）", min_value=8.0, max_value=72.0, value=18.0, step=0.5, key="h1_size")
            h1_bold = st.checkbox("加粗（H1）", value=True, key="h1_bold")
        with col_t2:
            st.markdown("**二级标题样式**")
            h2_font = st.text_input("字体名（H2）", value="SimHei", key="h2_font")
            h2_size = st.number_input("字号 pt（H2）", min_value=8.0, max_value=72.0, value=16.0, step=0.5, key="h2_size")
            h2_bold = st.checkbox("加粗（H2）", value=True, key="h2_bold")
        with col_t3:
            st.markdown("**三级标题样式**")
            h3_font = st.text_input("字体名（H3）", value="SimHei", key="h3_font")
            h3_size = st.number_input("字号 pt（H3）", min_value=8.0, max_value=72.0, value=14.0, step=0.5, key="h3_size")
            h3_bold = st.checkbox("加粗（H3）", value=False, key="h3_bold")

        if st.button("解析文档标题结构", disabled=not docx_title_file, key="btn_analyze_titles"):
            try:
                title_doc_bytes = read_uploaded_bytes(docx_title_file)
                rows = mp.analyze_docx_headings(title_doc_bytes)
                st.session_state["title_rows"] = rows
                c1 = sum(1 for r in rows if r.get("level") == 1)
                c2 = sum(1 for r in rows if r.get("level") == 2)
                c3 = sum(1 for r in rows if r.get("level") == 3)
                cn = sum(1 for r in rows if not r.get("level"))
                st.success(f"解析完成：H1={c1}, H2={c2}, H3={c3}, 普通段落={cn}")
            except Exception as e:
                st.error(f"解析失败：{type(e).__name__}: {e}")

        rows = st.session_state.get("title_rows", []) if docx_title_file else []
        if rows:
            st.dataframe(rows, use_container_width=True, height=320)

        if st.button("应用样式并导出", type="primary", disabled=not docx_title_file, key="btn_apply_title_style"):
            try:
                title_doc_bytes = read_uploaded_bytes(docx_title_file)
                cfg = {
                    1: {"font_name": h1_font, "size_pt": h1_size, "bold": h1_bold},
                    2: {"font_name": h2_font, "size_pt": h2_size, "bold": h2_bold},
                    3: {"font_name": h3_font, "size_pt": h3_size, "bold": h3_bold},
                }
                out_docx = mp.apply_title_formatting_to_docx(title_doc_bytes, cfg)
                st.success("标题样式已应用并生成新文档。")
                st.download_button(
                    "下载：TitleFormatted.docx",
                    data=out_docx,
                    file_name="TitleFormatted.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_title_formatted",
                )
            except Exception as e:
                st.error(f"导出失败：{type(e).__name__}: {e}")

        with st.expander("模板与样例测试说明", expanded=False):
            st.markdown(
                """
    **template.docx 如何构造**
    1. 在 Word 中新建文档。
    2. 为一级/二级/三级标题段落分别使用内置样式：`Heading 1`、`Heading 2`、`Heading 3`（中文界面常显示为“标题 1/2/3”）。
    3. 正文使用 `Normal`（正文）样式。
    4. 保存为 `template.docx` 后上传即可。

    **内置 sample 测试用例**
    - 点击下方按钮会在内存中构造一个 4 段文档：H1/H2/H3/正文各一段。
    - 程序会执行“解析 + 应用样式”，并检查计数与粗体设置是否符合预期。
                """
            )
            if st.button("运行内置 sample 测试", key="btn_title_selftest"):
                try:
                    result = mp.selftest_title_formatting()
                    if result.get("ok"):
                        st.success(f"Sample 测试通过：{result}")
                    else:
                        st.error(f"Sample 测试未通过：{result}")
                except Exception as e:
                    st.error(f"Sample 测试失败：{type(e).__name__}: {e}")
def docx_roundtrip_make_equations_editable(docx_bytes: bytes) -> bytes:
    sha = mp.sha256_bytes(docx_bytes)
    return _cached_roundtrip_omml(sha, docx_bytes)

def docx_roundtrip_make_equations_editable_with_md(docx_bytes: bytes):
    sha = mp.sha256_bytes(docx_bytes)
    return _cached_roundtrip_omml_with_md(sha, docx_bytes)

def docx_ai_roundtrip_make_equations_editable(
    docx_bytes: bytes,
    *,
    model: str,
    max_batch_chars: int,
    timeout_s: int,
    out_tokens: int,
    base_url: str,
):
    sha = mp.sha256_bytes(docx_bytes)
    return _cached_ai_roundtrip(
        sha, docx_bytes, model, int(max_batch_chars), int(timeout_s), int(out_tokens), str(base_url)
    )
