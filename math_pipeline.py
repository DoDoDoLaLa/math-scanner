# -*- coding: utf-8 -*-
"""
math_pipeline.py
集中放：DOCX<->Pandoc(Markdown/LaTeX)<->DOCX(OMML) 的“计算/转换”逻辑 + 统一的上传读取工具。

设计目标：
- app.py 尽量只做 UI（Streamlit 控件、布局、状态管理）。
- 这里提供纯函数（便于测试/缓存），不直接依赖 Streamlit。
"""
from __future__ import annotations

import io
import re
import json
import hashlib
import tempfile
from pathlib import Path
from typing import Callable, List, Optional, Tuple

import pypandoc
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


# -----------------------------
# Common helpers
# -----------------------------
def ensure_pandoc() -> None:
    """Ensure pandoc exists. In Streamlit Cloud, pypandoc-binary is recommended."""
    try:
        _ = pypandoc.get_pandoc_path()
    except OSError:
        try:
            pypandoc.download_pandoc()
        except Exception:
            # If download fails (e.g. no internet), rely on system/pypandoc-binary.
            pass


def sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def read_uploaded_bytes(uploaded_file) -> bytes:
    """Robustly read bytes from Streamlit UploadedFile (or file-like).

    Prefer .getvalue() to avoid consuming the internal pointer of the uploaded stream.
    """
    if uploaded_file is None:
        return b""
    if hasattr(uploaded_file, "getvalue"):
        return uploaded_file.getvalue()
    if hasattr(uploaded_file, "read"):
        return uploaded_file.read()
    raise TypeError(f"Unsupported uploaded object: {type(uploaded_file)}")


def normalize_md(md: str) -> str:
    md = (md or "").replace("\r\n", "\n").replace("\r", "\n")
    md = re.sub(r"\n{3,}", "\n\n", md).strip()
    return md


# -----------------------------
# Pandoc JSON AST semantic blocks (for robust chunking)
# -----------------------------
def pandoc_markdown_to_semantic_blocks(md: str) -> List[dict]:
    """Split markdown into top-level Pandoc blocks, returning a list of dicts:
    {type: <BlockType>, markdown: <block_markdown>}

    This is useful for translation chunking without breaking tables/lists/etc.
    """
    ensure_pandoc()
    md = normalize_md(md) + "\n"
    js = pypandoc.convert_text(md, to="json", format="markdown+tex_math_dollars")
    root = json.loads(js)
    api_ver = root.get("pandoc-api-version")
    blocks = root.get("blocks") or []
    out: List[dict] = []

    for b in blocks:
        b_type = (b.get("t") if isinstance(b, dict) else None) or "Unknown"
        mini = {"pandoc-api-version": api_ver, "meta": {}, "blocks": [b]}
        md_block = pypandoc.convert_text(json.dumps(mini, ensure_ascii=False), to="markdown+tex_math_dollars", format="json")
        out.append({"type": b_type, "markdown": normalize_md(md_block)})

    return out


def join_semantic_blocks(blocks: List[dict]) -> str:
    """Join blocks produced by pandoc_markdown_to_semantic_blocks back to markdown."""
    parts: List[str] = []
    for b in blocks or []:
        if isinstance(b, dict):
            parts.append((b.get("markdown") or "").strip())
        else:
            parts.append(str(b).strip())
    return normalize_md("\n\n".join([p for p in parts if p]))
# -----------------------------
# Pandoc math normalization / sanitization
# -----------------------------
PANDOC_DISPLAY_MATH_BRACKET_RE = re.compile(r"(?s)\\\[(.+?)\\\]")
PANDOC_DISPLAY_MATH_DOLLAR_RE = re.compile(r"(?s)\$\$(.+?)\$\$")
PANDOC_INLINE_MATH_PAREN_RE = re.compile(r"(?s)\\\((.+?)\\\)")
PANDOC_INLINE_MATH_DOLLAR_RE = re.compile(r"(?s)(?<!\$)\$([^$\n]+)\$(?!\$)")


def normalize_pandoc_math(
    text: str,
    *,
    display_style: str = "dollars",  # equation|equation*|bracket|dollars
    inline_style: str = "dollars",   # paren|dollars
) -> str:
    """Normalize delimiters so math is stable for later parsing and AI cleaning."""
    if not text:
        return text

    t = text.replace("\r\n", "\n").replace("\r", "\n")

    def _to_equation(body: str, starred: bool) -> str:
        env = "equation*" if starred else "equation"
        body = body.strip()
        return f"\\begin{{{env}}}\n{body}\n\\end{{{env}}}"

    def _display_repl(m):
        body = (m.group(1) or "").strip()
        if display_style == "equation":
            return _to_equation(body, starred=False)
        if display_style == "equation*":
            return _to_equation(body, starred=True)
        if display_style == "bracket":
            return f"\\[\n{body}\n\\]"
        # dollars
        return f"$$\n{body}\n$$"

    t = PANDOC_DISPLAY_MATH_BRACKET_RE.sub(_display_repl, t)
    t = PANDOC_DISPLAY_MATH_DOLLAR_RE.sub(_display_repl, t)

    def _inline_repl(m):
        body = (m.group(1) or "").strip()
        if inline_style == "paren":
            return f"\\({body}\\)"
        return f"${body}$"

    t = PANDOC_INLINE_MATH_PAREN_RE.sub(_inline_repl, t)
    t = PANDOC_INLINE_MATH_DOLLAR_RE.sub(lambda m: _inline_repl(m), t)
    return t


def sanitize_tex_math_for_pandoc(md: str) -> str:
    """Make TeX math more Pandoc-friendly:
    - Undo common escaping patterns (\$ -> $)
    - Ensure $$...$$ blocks contain NO blank lines (Pandoc can mis-parse)
    """
    if not md:
        return md

    t = md.replace("\r\n", "\n").replace("\r", "\n").replace("\u00A0", " ")
    t = t.replace(r"\\$", "$").replace(r"\$", "$")

    block_re = re.compile(r"(?s)\$\$(.+?)\$\$")
    inline_re = re.compile(r"(?s)(?<!\$)\$([^$\n]+)\$(?!\$)")

    def _clean_body(body: str, is_block: bool) -> str:
        b = (body or "").strip()
        if is_block:
            b = re.sub(r"\n\s*\n+", "\n", b)  # collapse blank lines
            b = re.sub(r"[ \t]+\n", "\n", b)
        else:
            b = b.replace("\n", " ")
            b = re.sub(r"\s{2,}", " ", b)
        return b.strip()

    def _fix_block(m):
        b = _clean_body(m.group(1), True)
        return "$$\n" + b + "\n$$"

    def _fix_inline(m):
        b = _clean_body(m.group(1), False)
        return "$" + b + "$"

    t = block_re.sub(_fix_block, t)
    t = inline_re.sub(_fix_inline, t)
    return t


# -----------------------------
# DOCX <-> Pandoc
# -----------------------------
def docx_to_pandoc_markdown_for_math(docx_bytes: bytes, *, wrap_none: bool = True) -> str:
    """DOCX -> Markdown(+tex_math_dollars) suitable for later math parsing."""
    if not docx_bytes:
        return ""

    ensure_pandoc()

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        in_path = td / "in.docx"
        in_path.write_bytes(docx_bytes)

        extra_args = ["--wrap=none"] if wrap_none else []
        md = pypandoc.convert_file(
            str(in_path),
            to="markdown+tex_math_dollars",
            format="docx",
            extra_args=extra_args,
        )

    md = (md or "").replace("\r\n", "\n").replace("\r", "\n")

    # Undo common escaping patterns that break TeX math parsing
    md = md.replace(r"\$", "$").replace(r"\\$", "$")

    # DOCX writer often escapes backslashes; be conservative:
    md = md.replace("\\\\", "\\")

    md = normalize_pandoc_math(md, display_style="dollars", inline_style="dollars")
    md = sanitize_tex_math_for_pandoc(md)
    return normalize_md(md)


def pandoc_markdown_to_docx(md: str, *, reference_docx_bytes: Optional[bytes] = None) -> bytes:
    """Markdown(+tex_math_dollars) -> DOCX. If reference_docx_bytes provided, preserve styles via --reference-doc."""
    ensure_pandoc()
    md = normalize_md(md) + "\n"

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        out_path = td / "out.docx"

        extra_args: List[str] = []
        if reference_docx_bytes:
            ref = td / "ref.docx"
            ref.write_bytes(reference_docx_bytes)
            extra_args += ["--reference-doc", str(ref)]

        pypandoc.convert_text(
            md,
            to="docx",
            format="markdown+tex_math_dollars",
            outputfile=str(out_path),
            extra_args=extra_args,
        )
        return out_path.read_bytes()




def pandoc_markdown_to_latex(md: str) -> str:
    """Markdown(+tex_math_dollars) -> LaTeX string (for export)."""
    ensure_pandoc()
    md = normalize_md(md) + "\n"
    return pypandoc.convert_text(md, to="latex", format="markdown+tex_math_dollars")


def docx_roundtrip_make_equations_editable(docx_bytes: bytes) -> Tuple[bytes, str]:
    """DOCX -> MD(tex_math_dollars) -> sanitize -> DOCX (Pandoc emits OMML). Returns (out_docx, md_used)."""
    md = docx_to_pandoc_markdown_for_math(docx_bytes, wrap_none=True)
    out = pandoc_markdown_to_docx(md, reference_docx_bytes=docx_bytes)
    return out, md


def extract_docx_preview_text_for_glossary(docx_bytes: bytes, *, max_chars: int = 9000) -> str:
    """Extract a lightweight text preview from DOCX for smart glossary mining.

    Uses the existing Pandoc markdown pipeline so equations and document ordering remain
    relatively stable, then returns a clipped plain-text-like preview.
    """
    if not docx_bytes:
        return ""
    md = docx_to_pandoc_markdown_for_math(docx_bytes, wrap_none=True)
    if not md:
        return ""
    t = md.replace("\r\n", "\n").replace("\r", "\n")
    t = re.sub(r"```[\s\S]*?```", "\n", t)
    t = re.sub(r"\$\$[\s\S]*?\$\$", " [MATH_BLOCK] ", t)
    t = re.sub(r"(?<!\$)\$([^$\n]+)\$(?!\$)", r" [MATH_INLINE: \1] ", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t[: int(max_chars)]




# -----------------------------
# Title formatting (python-docx only, no AI)
# -----------------------------
def _extract_heading_level_from_style_name(style_name: str) -> Optional[int]:
    if not style_name:
        return None
    name = str(style_name).strip()
    m = re.search(r"(?:^|\s)Heading\s*([1-9])(?:\s|$)", name, flags=re.IGNORECASE)
    if m:
        return int(m.group(1))
    m = re.search(r"(?:^|\s)标题\s*([1-9])(?:\s|$)", name)
    if m:
        return int(m.group(1))
    if re.search(r"heading\s*1", name, flags=re.IGNORECASE) or "标题1" in name:
        return 1
    if re.search(r"heading\s*2", name, flags=re.IGNORECASE) or "标题2" in name:
        return 2
    if re.search(r"heading\s*3", name, flags=re.IGNORECASE) or "标题3" in name:
        return 3
    return None


def detect_heading_level_for_paragraph(paragraph) -> Optional[int]:
    """Detect heading level by paragraph style only (no AI).

    Returns 1/2/3 for title levels, otherwise None.
    """
    try:
        style_name = paragraph.style.name if paragraph.style is not None else ""
    except Exception:
        style_name = ""
    lv = _extract_heading_level_from_style_name(style_name)
    if lv in (1, 2, 3):
        return lv
    return None


def analyze_docx_headings(docx_bytes: bytes) -> List[dict]:
    """Analyze each paragraph and classify as heading1/2/3 or normal."""
    if not docx_bytes:
        return []
    doc = Document(io.BytesIO(docx_bytes))
    out: List[dict] = []
    for idx, p in enumerate(doc.paragraphs, start=1):
        lv = detect_heading_level_for_paragraph(p)
        ptype = f"Heading {lv}" if lv else "Normal"
        out.append({
            "index": idx,
            "text": (p.text or "").strip(),
            "type": ptype,
            "level": lv,
            "style_name": p.style.name if p.style is not None else "",
        })
    return out


def _set_run_font(run, *, font_name: str, size_pt: float, bold: bool) -> None:
    run.font.name = font_name
    run.bold = bool(bold)
    run.font.bold = bool(bold)
    run.font.size = Pt(float(size_pt))
    # Ensure East Asian font also follows configured name.
    try:
        rfonts = run._element.rPr.rFonts if run._element.rPr is not None else None
        if rfonts is None:
            run._element.get_or_add_rPr().get_or_add_rFonts()
            rfonts = run._element.rPr.rFonts
        rfonts.set(qn("w:eastAsia"), font_name)
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
    except Exception:
        pass


def apply_title_formatting_to_docx(
    docx_bytes: bytes,
    level_styles: dict,
    *,
    apply_to_body: bool = False,
    body_style: Optional[dict] = None,
    apply_to_tables: bool = False,
    table_style: Optional[dict] = None,
    apply_to_captions: bool = False,
    caption_style: Optional[dict] = None,
) -> bytes:
    """Apply user-configured run font formatting.

    Backward compatible:
    - Always formats Heading 1/2/3 based on detect_heading_level_for_paragraph.
    - Optionally (when toggled) also formats:
        * body paragraphs (non-heading)
        * tables (all cell paragraphs)
        * captions (heuristic by style name)
    """
    if not docx_bytes:
        return docx_bytes

    doc = Document(io.BytesIO(docx_bytes))

    # Prepare optional styles (safe defaults)
    body_style = body_style or {}
    table_style = table_style or {}
    caption_style = caption_style or {}

    def _norm_style(cfg: dict, default_font: str, default_size: float, default_bold: bool) -> Tuple[str, float, bool]:
        font_name = str(cfg.get("font_name") or default_font)
        size_pt = float(cfg.get("size_pt") or default_size)
        bold = bool(cfg.get("bold") if cfg.get("bold") is not None else default_bold)
        return font_name, size_pt, bold

    def _is_caption_para(p) -> bool:
        try:
            s = (p.style.name or "").lower()
        except Exception:
            s = ""
        if not s:
            return False
        # Common caption styles in Word / Chinese Word
        keys = ["caption", "题注", "图题", "表题", "figure caption", "table caption"]
        return any(k.lower() in s for k in keys)

    # 1) headings (always on)
    for p in doc.paragraphs:
        lv = detect_heading_level_for_paragraph(p)
        if lv not in (1, 2, 3):
            continue
        cfg = level_styles.get(int(lv)) or {}
        font_name, size_pt, bold = _norm_style(cfg, "SimHei", 16.0, True)
        for r in (p.runs or []):
            _set_run_font(r, font_name=font_name, size_pt=size_pt, bold=bold)

    # 2) body/captions (optional)
    if apply_to_body or apply_to_captions:
        b_font, b_size, b_bold = _norm_style(body_style, "Times New Roman", 11.0, False)
        c_font, c_size, c_bold = _norm_style(caption_style, b_font, b_size, b_bold)

        for p in doc.paragraphs:
            lv = detect_heading_level_for_paragraph(p)
            if lv in (1, 2, 3):
                continue

            if apply_to_captions and _is_caption_para(p):
                font_name, size_pt, bold = c_font, c_size, c_bold
            elif apply_to_body:
                font_name, size_pt, bold = b_font, b_size, b_bold
            else:
                continue

            if not p.runs:
                p.add_run(p.text or "")
            for r in p.runs:
                _set_run_font(r, font_name=font_name, size_pt=size_pt, bold=bold)

    # 3) tables (optional)
    if apply_to_tables:
        t_font, t_size, t_bold = _norm_style(table_style, "Times New Roman", 10.5, False)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if not p.runs:
                            p.add_run(p.text or "")
                        for r in p.runs:
                            _set_run_font(r, font_name=t_font, size_pt=t_size, bold=t_bold)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def selftest_title_formatting() -> dict:
    """Built-in sample test case for heading analysis + formatting."""
    d = Document()
    p1 = d.add_paragraph("第一章 介绍", style="Heading 1")
    p2 = d.add_paragraph("1.1 背景", style="Heading 2")
    p3 = d.add_paragraph("1.1.1 定义", style="Heading 3")
    p4 = d.add_paragraph("普通正文段落", style="Normal")
    b = io.BytesIO(); d.save(b)

    analyzed = analyze_docx_headings(b.getvalue())
    counts = {"Heading 1": 0, "Heading 2": 0, "Heading 3": 0, "Normal": 0}
    for x in analyzed:
        counts[x["type"]] = counts.get(x["type"], 0) + 1

    styled = apply_title_formatting_to_docx(
        b.getvalue(),
        {
            1: {"font_name": "SimHei", "size_pt": 18, "bold": True},
            2: {"font_name": "SimHei", "size_pt": 16, "bold": True},
            3: {"font_name": "SimHei", "size_pt": 14, "bold": False},
        },
    )
    d2 = Document(io.BytesIO(styled))
    h1 = d2.paragraphs[0].runs[0]
    h3 = d2.paragraphs[2].runs[0]

    ok = (
        counts.get("Heading 1", 0) == 1
        and counts.get("Heading 2", 0) == 1
        and counts.get("Heading 3", 0) == 1
        and counts.get("Normal", 0) == 1
        and bool(h1.bold) is True
        and bool(h3.bold) is False
    )
    return {"ok": bool(ok), "counts": counts}

# -----------------------------
# AI-assisted cleaning (app passes in a callable)
# -----------------------------
AI_MATH_CLEAN_PROMPT_ZH = """你是论文排版助理。请把我给你的 Markdown 文本“尽量保持原文不变”，只对其中的数学公式做纠错与统一格式，使其更像“图片 OCR 输出的学术 LaTeX”风格，便于后续用 Pandoc 解析成 Word 可编辑公式（OMML）。

严格要求：
1) 仅处理数学公式：位于 $...$ 或 $$...$$ 或 \\( ... \\) 或 \\[ ... \\] 或 \\begin{equation}...\\end{equation} 内的内容。普通文本不要改写措辞。
2) 统一分隔符：
   - 行内公式统一成 $...$
   - 行间公式统一成 $$...$$（注意：$$...$$ 内部不要出现空行；如果有换行，保持为单个换行即可）
3) 纠错策略（尽量保守，宁可不改也不要瞎改）：
   - 把明显的“文本型下标/上标”统一为直立体：例如 P_seed -> P_{\\mathrm{seed}}；t_k^+ 中的 + 保留；类似 \\Delta_h、\\rho_h 这类命令保持不变。
   - 保留 \\alpha \\beta 等 LaTeX 命令；不要把 \\\\alpha 之类的转义搞坏。
   - 不要引入新的宏包命令；不要新增 \\label/\\ref 等。
4) 不要输出解释，不要输出代码块围栏，不要输出 JSON。只输出修正后的 Markdown 纯文本。
"""


def chunk_text_by_paragraph(md: str, *, max_chars: int) -> List[str]:
    if not md:
        return []
    paras = md.split("\n\n")
    chunks: List[str] = []
    cur: List[str] = []
    cur_len = 0
    for p in paras:
        p2 = p.strip("\n")
        if not p2:
            continue
        add = p2 + "\n\n"
        if cur and cur_len + len(add) > max_chars:
            chunks.append("".join(cur).strip() + "\n")
            cur, cur_len = [], 0
        cur.append(add)
        cur_len += len(add)
    if cur:
        chunks.append("".join(cur).strip() + "\n")
    return chunks


def ai_clean_markdown_math_like_ocr(
    md: str,
    *,
    call_llm: Callable[[str], str],
    max_batch_chars: int,
    progress_cb: Optional[Callable[[int, int], None]] = None,
) -> str:
    """LLM cleans markdown math. call_llm(prompt_text)->response_text should be provided by app."""
    md = normalize_md(md)
    if not md:
        return md

    parts = chunk_text_by_paragraph(md, max_chars=int(max_batch_chars))
    out_parts: List[str] = []
    total = len(parts)

    for i, part in enumerate(parts, start=1):
        if progress_cb:
            progress_cb(i, total)
        prompt = AI_MATH_CLEAN_PROMPT_ZH.strip() + "\n\n---\n\n" + part
        cleaned = normalize_md(call_llm(prompt))
        out_parts.append(cleaned)

    merged = normalize_md("\n\n".join([p for p in out_parts if p.strip()]))
    merged = normalize_pandoc_math(merged, display_style="dollars", inline_style="dollars")
    merged = sanitize_tex_math_for_pandoc(merged)
    return normalize_md(merged)


def docx_ai_roundtrip_make_equations_editable(
    docx_bytes: bytes,
    *,
    call_llm: Callable[[str], str],
    max_batch_chars: int,
    progress_cb: Optional[Callable[[int, int], None]] = None,
) -> Tuple[bytes, str]:
    """DOCX -> MD -> AI clean -> DOCX(OMML). Returns (out_docx, cleaned_md)."""
    md = docx_to_pandoc_markdown_for_math(docx_bytes, wrap_none=True)
    md2 = ai_clean_markdown_math_like_ocr(
        md,
        call_llm=call_llm,
        max_batch_chars=int(max_batch_chars),
        progress_cb=progress_cb,
    )
    out = pandoc_markdown_to_docx(md2, reference_docx_bytes=docx_bytes)
    return out, md2
