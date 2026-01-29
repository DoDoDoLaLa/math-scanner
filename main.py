import streamlit as st
import google.generativeai as genai
from PIL import Image
import io
import matplotlib.pyplot as plt
import matplotlib as mpl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import time

# ==========================================
# 0. 全局配置
# ==========================================
st.set_page_config(
    page_title="MathScan Pro - 印刷级公式还原",
    page_icon="✒️",
    layout="wide"
)

# 字体优化：尝试使用标准数学字体
try:
    mpl.rcParams['mathtext.fontset'] = 'cm'
    mpl.rcParams['font.family'] = 'serif'
except:
    pass

# ==========================================
# 1. 智能 AI 识别模块 (自动降级与轮询)
# ==========================================
def extract_content_with_gemini(api_key, image):
    # 配置 API
    try:
        genai.configure(api_key=api_key)
    except Exception as e:
        st.error(f"API Key 配置失败: {e}")
        return None

    # 定义模型尝试列表：优先使用最强模型，如果环境不支持则自动降级
    # 1. gemini-1.5-pro: 最强，识别复杂公式下标最准
    # 2. gemini-1.5-flash: 速度最快
    # 3. gemini-pro-vision: 1.0版本，兼容旧版 SDK (防止 404 错误)
    model_candidates = [
        "gemini-1.5-pro",
        "gemini-1.5-flash",
        "gemini-pro-vision"
    ]

    prompt = """
    你是一个专业的排版转换工具。请将图片内容转换为文档，要求如下：
    1. 【识别精准】：识别图片中的所有中文、英文文本和数学公式。
    2. 【LaTeX 格式】：所有数学公式必须写成标准的 LaTeX 格式。
       - 简单的变量（如 x, N, t）用单美元符号：$x$
       - 复杂的公式（如带上下标、求和、分数）用双美元符号：$$ ... $$
    3. 【保持结构】：严格保持原文的段落结构。
    4. 【输出纯净】：只输出内容，不要任何 Markdown 代码块标记，不要 "Here is the output" 等废话。
    """

    last_error = None
    
    # 创建一个空的占位符用于显示正在尝试的状态
    status_placeholder = st.empty()

    for model_name in model_candidates:
        try:
            status_placeholder.info(f"🔄 正在尝试连接模型: **{model_name}** ...")
            
            # 针对不同模型的配置微调
            generation_config = {"temperature": 0.1, "max_output_tokens": 4096}
            model = genai.GenerativeModel(model_name=model_name, generation_config=generation_config)
            
            # 发送请求
            response = model.generate_content([prompt, image])
            text = response.text
            
            # 清理 Markdown 标记
            text = text.replace("```latex", "").replace("```markdown", "").replace("```", "")
            
            status_placeholder.success(f"✅ 成功连接模型: {model_name}")
            time.sleep(1) # 让用户看到成功提示
            status_placeholder.empty() # 清除提示
            return text.strip()

        except Exception as e:
            error_str = str(e)
            print(f"模型 {model_name} 失败: {error_str}")
            # 如果是 404 (模型未找到)，继续尝试下一个；如果是 403 (Key无效)，直接停止
            if "403" in error_str:
                status_placeholder.error("API Key 无效或没有权限。请检查 Key 是否正确。")
                return None
            last_error = e
            continue
    
    # 如果所有模型都失败了
    status_placeholder.empty()
    st.error("❌ 所有模型连接均失败。")
    
    # 给出具体的调试建议
    if last_error:
        err_msg = str(last_error)
        if "404" in err_msg:
            st.error("⚠️ 核心原因：服务器 SDK 版本过旧，找不到新模型。")
            st.warning("👉 请务必在 Streamlit 后台点击 'Reboot' (重启) 以强制更新环境。")
            st.code(f"当前检测到的 SDK 版本: {genai.__version__}\n(需要 >= 0.7.2 才能支持 1.5 Pro)", language="text")
        else:
            st.error(f"错误详情: {err_msg}")
            
    return None

# ==========================================
# 2. 公式渲染模块 (转透明高清图)
# ==========================================
def render_latex_high_quality(latex_str, dpi=400):
    try:
        latex_str = latex_str.strip()
        # 创建画布
        fig = plt.figure(figsize=(0.1, 0.1))
        # 渲染内容
        text_content = f"${latex_str}$"
        # 绘制黑色文字
        fig.text(0.5, 0.5, text_content, fontsize=20, ha='center', va='center', color='black')
        plt.axis('off')
        
        # 保存到内存
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=dpi, bbox_inches='tight', pad_inches=0.1, transparent=True)
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception as e:
        # 渲染失败时不打印到前台干扰用户，后台记录即可
        print(f"Render Error: {e}")
        plt.close(fig)
        return None

# ==========================================
# 3. 文档生成模块
# ==========================================
def create_professional_doc(raw_text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    
    # 标题
    doc.add_heading('MathScan 识别结果', 0)
    
    # 拆分文本和公式块 ($$)
    parts = re.split(r'(\$\$[\s\S]*?\$\$)', raw_text)
    
    for part in parts:
        part = part.strip()
        if not part:
            continue
            
        if part.startswith('$$') and part.endswith('$$'):
            # === 处理独立公式 ===
            latex_code = part[2:-2].strip()
            img_stream = render_latex_high_quality(latex_code)
            
            if img_stream:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                # 插入图片，高度 0.6 英寸 (约 1.5cm)，既清晰又不会太大
                run.add_picture(img_stream, height=Inches(0.6)) 
            else:
                # 失败兜底
                p = doc.add_paragraph(latex_code)
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        else:
            # === 处理普通文本 ===
            doc.add_paragraph(part)
            
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ==========================================
# 4. 界面主逻辑
# ==========================================
st.title("✒️ MathScan Pro：印刷级公式还原 (智能版)")

# 诊断信息：在侧边栏底部悄悄显示版本，方便调试
with st.sidebar:
    st.divider()
    st.caption(f"🔧 SDK Version: {genai.__version__}")
    st.caption("如果是 0.5.x 或更低，请更新 requirements.txt 并重启 App")

# 输入框置顶
with st.container():
    st.info("💡 请输入您的 Google API Key (以 AIza 开头)")
    api_key = st.text_input("API Key", type="password", label_visibility="collapsed", placeholder="粘贴 Key 到这里...")

st.markdown("---")
uploaded_file = st.file_uploader("上传图片", type=["png", "jpg", "jpeg"])

if uploaded_file and api_key:
    image = Image.open(uploaded_file)
    st.image(image, caption="预览", width=400)
    
    if st.button("开始转换 / Generate", type="primary"):
        text_result = extract_content_with_gemini(api_key, image)
        
        if text_result:
            st.success("识别成功！正在生成文档...")
            
            # 显示文本预览
            with st.expander("查看识别结果 (LaTeX 源码)"):
                st.code(text_result, language='latex')
            
            # 生成文档
            doc_file = create_professional_doc(text_result)
            
            st.download_button(
                label="📥 下载 Word 文档 (.docx)",
                data=doc_file,
                file_name="公式还原结果.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

elif uploaded_file and not api_key:
    st.warning("请在上方输入 API Key 才能开始运行。")